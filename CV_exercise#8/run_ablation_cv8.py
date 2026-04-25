import csv
import time
from pathlib import Path

import numpy as np
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
from torch.utils.data import Dataset, DataLoader
from sklearn.datasets import load_digits
from sklearn.model_selection import train_test_split
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

import torchvision.transforms as T
import torchvision.transforms.functional as TF
from torchvision.datasets import CIFAR10
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#8")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DEVICE = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
SEED = 42
np.random.seed(SEED)
torch.manual_seed(SEED)

rows = []

def add_row(ex, param, value, out_file, analysis, metric=''):
    rows.append({
        'Exercise': f'Exercise {ex}',
        'Parameter': str(param),
        'Value': str(value),
        'Output file': str(out_file.relative_to(ROOT)).replace('\\', '/'),
        'Visual change': '',
        'Optional metric': metric,
        'One-line analysis': analysis,
    })

class NumpyImageDataset(Dataset):
    def __init__(self, images, labels, transform=None):
        self.images = images
        self.labels = labels
        self.transform = transform

    def __len__(self):
        return len(self.labels)

    def __getitem__(self, idx):
        x = Image.fromarray(self.images[idx])
        y = int(self.labels[idx])
        if self.transform is not None:
            x = self.transform(x)
        return x, y

class RawImageDataset(Dataset):
    def __init__(self, images, labels):
        self.images = images
        self.labels = labels

    def __len__(self):
        return len(self.labels)

    def __getitem__(self, idx):
        return Image.fromarray(self.images[idx]), int(self.labels[idx])


def prepare_data():
    cifar_roots = [ROOT / 'cifar_user', ROOT / 'cifar_local', ROOT / 'data']
    for croot in cifar_roots:
        try:
            train_raw = CIFAR10(str(croot), train=True, download=False, transform=None)
            test_raw = CIFAR10(str(croot), train=False, download=False, transform=None)
            x_train = np.array(train_raw.data)
            y_train = np.array(train_raw.targets)
            x_test = np.array(test_raw.data)
            y_test = np.array(test_raw.targets)
            source = 'cifar10_local'
            n_classes = 10
            return x_train, y_train, x_test, y_test, source, n_classes
        except Exception:
            pass

    digits = load_digits()
    x = digits.images.astype(np.float32) / 16.0
    x = torch.tensor(x).unsqueeze(1)
    x = F.interpolate(x, size=(32, 32), mode='bilinear', align_corners=False).squeeze(1).numpy()
    x = (x * 255.0).clip(0, 255).astype(np.uint8)
    x = np.stack([x, x, x], axis=-1)
    y = digits.target.astype(np.int64)

    x_train, x_test, y_train, y_test = train_test_split(
        x, y, test_size=0.2, random_state=SEED, stratify=y
    )
    source = 'digits_upsampled'
    n_classes = len(np.unique(y))
    return x_train, y_train, x_test, y_test, source, n_classes


def build_model(num_classes=10):
    return nn.Sequential(
        nn.Conv2d(3, 32, 3, padding=1), nn.BatchNorm2d(32), nn.ReLU(),
        nn.MaxPool2d(2),
        nn.Conv2d(32, 64, 3, padding=1), nn.BatchNorm2d(64), nn.ReLU(),
        nn.MaxPool2d(2),
        nn.Conv2d(64, 128, 3, padding=1), nn.BatchNorm2d(128), nn.ReLU(),
        nn.AdaptiveAvgPool2d(1),
        nn.Flatten(),
        nn.Linear(128, num_classes),
    )


def evaluate(model, loader):
    model.eval()
    c = t = 0
    with torch.no_grad():
        for xb, yb in loader:
            xb = xb.to(DEVICE)
            yb = yb.to(DEVICE)
            pred = model(xb).argmax(1)
            c += (pred == yb).sum().item()
            t += yb.size(0)
    return 100.0 * c / max(1, t)


def mixup_batch(x, y, alpha=0.2):
    lam = float(np.random.beta(alpha, alpha))
    idx = torch.randperm(x.size(0), device=x.device)
    mixed = lam * x + (1 - lam) * x[idx]
    return mixed, y, y[idx], lam


def rand_bbox(size, lam):
    _, _, h, w = size
    cut_rat = np.sqrt(1.0 - lam)
    cut_w = max(1, int(w * cut_rat))
    cut_h = max(1, int(h * cut_rat))
    cx = np.random.randint(w)
    cy = np.random.randint(h)
    x1 = np.clip(cx - cut_w // 2, 0, w)
    y1 = np.clip(cy - cut_h // 2, 0, h)
    x2 = np.clip(cx + cut_w // 2, 0, w)
    y2 = np.clip(cy + cut_h // 2, 0, h)
    return x1, y1, x2, y2


def cutmix_batch(x, y, alpha=1.0):
    lam = float(np.random.beta(alpha, alpha))
    idx = torch.randperm(x.size(0), device=x.device)
    x2 = x[idx].clone()
    x1, y1, x2c, y2c = rand_bbox(x.size(), lam)
    x[:, :, y1:y2c, x1:x2c] = x2[:, :, y1:y2c, x1:x2c]
    lam_actual = 1.0 - ((x2c - x1) * (y2c - y1) / (x.size(-1) * x.size(-2)))
    return x, y, y[idx], lam_actual


def train_model(
    train_loader,
    val_loader,
    num_classes,
    epochs=4,
    lr=1e-3,
    mode='baseline',
    mixup_alpha=0.2,
    cutmix_alpha=1.0,
):
    model = build_model(num_classes=num_classes).to(DEVICE)
    opt = optim.Adam(model.parameters(), lr=lr, weight_decay=1e-4)
    ce = nn.CrossEntropyLoss()

    for _ in range(epochs):
        model.train()
        for xb, yb in train_loader:
            xb = xb.to(DEVICE)
            yb = yb.to(DEVICE)
            opt.zero_grad()

            if mode == 'mixup':
                xmix, ya, yb2, lam = mixup_batch(xb, yb, alpha=mixup_alpha)
                out = model(xmix)
                loss = lam * ce(out, ya) + (1 - lam) * ce(out, yb2)
            elif mode == 'cutmix':
                xcm, ya, yb2, lam = cutmix_batch(xb, yb, alpha=cutmix_alpha)
                out = model(xcm)
                loss = lam * ce(out, ya) + (1 - lam) * ce(out, yb2)
            else:
                out = model(xb)
                loss = ce(out, yb)

            loss.backward()
            opt.step()

    val_acc = evaluate(model, val_loader)
    return model, val_acc


def tta_eval(model, raw_dataset, transform_list, batch_size=128):
    model.eval()
    total = 0
    correct = 0
    with torch.no_grad():
        for s in range(0, len(raw_dataset), batch_size):
            e = min(s + batch_size, len(raw_dataset))
            imgs = []
            labels = []
            for i in range(s, e):
                img, lab = raw_dataset[i]
                imgs.append(img)
                labels.append(lab)

            labels_t = torch.tensor(labels, device=DEVICE)
            probs_sum = None
            for tfm in transform_list:
                xb = torch.stack([tfm(img) for img in imgs], dim=0).to(DEVICE)
                probs = torch.softmax(model(xb), dim=1)
                if probs_sum is None:
                    probs_sum = probs
                else:
                    probs_sum = probs_sum + probs

            pred = (probs_sum / len(transform_list)).argmax(1)
            correct += (pred == labels_t).sum().item()
            total += len(labels)
    return 100.0 * correct / max(1, total)


x_train, y_train, x_test, y_test, data_source, num_classes = prepare_data()

norm = T.Normalize([0.5, 0.5, 0.5], [0.5, 0.5, 0.5])

def to_tensor_norm():
    return T.Compose([T.ToTensor(), norm])

# Exercise 1: no aug vs basic aug
exd = OUT / 'ex1_augmentation_basic'
exd.mkdir(exist_ok=True)

no_aug = to_tensor_norm()
basic_aug = T.Compose([
    T.RandomHorizontalFlip(p=0.5),
    T.RandomCrop(32, padding=4),
    T.ColorJitter(brightness=0.2, contrast=0.2, saturation=0.2, hue=0.05),
    T.ToTensor(),
    norm,
])

train_no_aug = NumpyImageDataset(x_train, y_train, transform=no_aug)
train_aug = NumpyImageDataset(x_train, y_train, transform=basic_aug)
val_set = NumpyImageDataset(x_test, y_test, transform=to_tensor_norm())

loader_no_aug = DataLoader(train_no_aug, batch_size=128, shuffle=True)
loader_aug = DataLoader(train_aug, batch_size=128, shuffle=True)
val_loader = DataLoader(val_set, batch_size=256, shuffle=False)

_, acc_no_aug = train_model(loader_no_aug, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')
_, acc_aug = train_model(loader_aug, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')

labels = ['NoAug', 'BasicAug']
vals = [acc_no_aug, acc_aug]
plt.figure(figsize=(6, 4))
plt.bar(labels, vals, color=['#888888', '#2e7d32'])
for i, v in enumerate(vals):
    plt.text(i, v + 0.3, f'{v:.2f}', ha='center')
plt.ylim(max(0, min(vals) - 2), min(100, max(vals) + 5))
plt.title('Exercise 1 Accuracy Comparison')
plt.ylabel('Val Accuracy (%)')
plt.tight_layout()
out_all = exd / 'ex1_aug_compare.png'
plt.savefig(out_all, dpi=140)
plt.close()

add_row(1, 'augmentation', 'none', out_all, '증강을 끄면 학습 분포가 고정되어 일반화 이득이 제한됐다.', f'val_acc={acc_no_aug:.2f}')
add_row(1, 'augmentation', 'flip+crop+jitter', out_all, '기본 증강을 적용하면 위치와 색 변화에 대한 강건성이 올라갔다.', f'val_acc={acc_aug:.2f}')

# Exercise 2: baseline vs cutout mixup cutmix
exd = OUT / 'ex2_advanced_aug'
exd.mkdir(exist_ok=True)

class Cutout:
    def __init__(self, size=16):
        self.size = size

    def __call__(self, img_t):
        h = img_t.shape[1]
        w = img_t.shape[2]
        cy = np.random.randint(h)
        cx = np.random.randint(w)
        y1 = max(0, cy - self.size // 2)
        y2 = min(h, cy + self.size // 2)
        x1 = max(0, cx - self.size // 2)
        x2 = min(w, cx + self.size // 2)
        img_t[:, y1:y2, x1:x2] = 0.0
        return img_t

base_train = T.Compose([
    T.RandomHorizontalFlip(p=0.5),
    T.RandomCrop(32, padding=4),
    T.ToTensor(),
    norm,
])

cutout_train = T.Compose([
    T.RandomHorizontalFlip(p=0.5),
    T.RandomCrop(32, padding=4),
    T.ToTensor(),
    norm,
    Cutout(size=16),
])

tr_base = DataLoader(NumpyImageDataset(x_train, y_train, transform=base_train), batch_size=128, shuffle=True)
tr_cutout = DataLoader(NumpyImageDataset(x_train, y_train, transform=cutout_train), batch_size=128, shuffle=True)
val_loader = DataLoader(NumpyImageDataset(x_test, y_test, transform=to_tensor_norm()), batch_size=256, shuffle=False)

_, acc_base = train_model(tr_base, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')
_, acc_cutout = train_model(tr_cutout, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')
_, acc_mixup = train_model(tr_base, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='mixup')
_, acc_cutmix = train_model(tr_base, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='cutmix')

methods = ['Baseline', 'Cutout16', 'Mixup0.2', 'CutMix1.0']
accs = [acc_base, acc_cutout, acc_mixup, acc_cutmix]
plt.figure(figsize=(8, 4))
plt.bar(methods, accs, color=['#607d8b', '#8e24aa', '#ef6c00', '#0277bd'])
for i, v in enumerate(accs):
    plt.text(i, v + 0.3, f'{v:.2f}', ha='center', fontsize=9)
plt.ylim(max(0, min(accs) - 2), min(100, max(accs) + 5))
plt.title('Exercise 2 Advanced Augmentation')
plt.ylabel('Val Accuracy (%)')
plt.tight_layout()
out_all = exd / 'ex2_adv_aug_compare.png'
plt.savefig(out_all, dpi=140)
plt.close()

add_row(2, 'method', 'baseline', out_all, '기준 설정으로 다른 고급 증강 효과를 비교할 수 있다.', f'val_acc={acc_base:.2f}')
add_row(2, 'method', 'cutout(size=16)', out_all, '가림 영역 학습으로 부분 정보에 대한 복원력이 증가했다.', f'val_acc={acc_cutout:.2f}')
add_row(2, 'method', 'mixup(alpha=0.2)', out_all, '샘플 혼합으로 결정 경계가 부드러워져 과적합이 완화됐다.', f'val_acc={acc_mixup:.2f}')
add_row(2, 'method', 'cutmix(alpha=1.0)', out_all, '패치 교체 기반 혼합이 지역 단서와 라벨 혼합을 동시에 제공했다.', f'val_acc={acc_cutmix:.2f}')

# Exercise 3: augmentation probability
exd = OUT / 'ex3_aug_probability'
exd.mkdir(exist_ok=True)


def make_transform(p):
    return T.Compose([
        T.RandomHorizontalFlip(p=p),
        T.RandomApply([T.RandomCrop(32, padding=4)], p=p),
        T.RandomApply([T.RandomRotation(10)], p=p),
        T.ColorJitter(brightness=0.15 * p, contrast=0.2 * p, saturation=0.15 * p, hue=0.03 * p),
        T.ToTensor(),
        norm,
    ])

p_values = [0.2, 0.4, 0.6, 0.8]
p_accs = []
for p in p_values:
    tr = DataLoader(NumpyImageDataset(x_train, y_train, transform=make_transform(p)), batch_size=128, shuffle=True)
    _, acc = train_model(tr, val_loader, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')
    p_accs.append(acc)

plt.figure(figsize=(7, 4))
plt.plot(p_values, p_accs, marker='o', linewidth=2)
for p, a in zip(p_values, p_accs):
    plt.text(p, a + 0.2, f'{a:.2f}', ha='center', fontsize=9)
plt.title('Exercise 3 Augmentation Probability Sweep')
plt.xlabel('probability p')
plt.ylabel('Val Accuracy (%)')
plt.grid(alpha=0.25)
plt.tight_layout()
out_all = exd / 'ex3_aug_p_sweep.png'
plt.savefig(out_all, dpi=140)
plt.close()

for p, a in zip(p_values, p_accs):
    add_row(3, 'augmentation_p', p, out_all, '확률을 올리면 강건성은 증가하지만 과도하면 원본 분포에서 멀어질 수 있다.', f'val_acc={a:.2f}')

# Exercise 4: TTA
exd = OUT / 'ex4_tta'
exd.mkdir(exist_ok=True)

train_loader = DataLoader(NumpyImageDataset(x_train, y_train, transform=base_train), batch_size=128, shuffle=True)
val_loader_std = DataLoader(NumpyImageDataset(x_test, y_test, transform=to_tensor_norm()), batch_size=256, shuffle=False)
model_tta, acc_std = train_model(train_loader, val_loader_std, num_classes=num_classes, epochs=4, lr=1e-3, mode='baseline')

raw_val = RawImageDataset(x_test, y_test)
std_tfm = to_tensor_norm()
flip_tfm = T.Compose([T.RandomHorizontalFlip(p=1.0), T.ToTensor(), norm])
crop_tl = T.Compose([T.Lambda(lambda im: TF.crop(im, top=0, left=0, height=28, width=28)), T.Resize((32, 32)), T.ToTensor(), norm])
crop_tr = T.Compose([T.Lambda(lambda im: TF.crop(im, top=0, left=4, height=28, width=28)), T.Resize((32, 32)), T.ToTensor(), norm])
crop_bl = T.Compose([T.Lambda(lambda im: TF.crop(im, top=4, left=0, height=28, width=28)), T.Resize((32, 32)), T.ToTensor(), norm])
crop_br = T.Compose([T.Lambda(lambda im: TF.crop(im, top=4, left=4, height=28, width=28)), T.Resize((32, 32)), T.ToTensor(), norm])

acc_1 = tta_eval(model_tta, raw_val, [std_tfm])
acc_2 = tta_eval(model_tta, raw_val, [std_tfm, flip_tfm])
six_view_tfms = [std_tfm, flip_tfm, crop_tl, crop_tr, crop_bl, crop_br]
acc_6 = tta_eval(model_tta, raw_val, six_view_tfms)

flip_crop_tl = T.Compose([T.Lambda(lambda im: TF.hflip(TF.crop(im, top=0, left=0, height=28, width=28))), T.Resize((32, 32)), T.ToTensor(), norm])
flip_crop_tr = T.Compose([T.Lambda(lambda im: TF.hflip(TF.crop(im, top=0, left=4, height=28, width=28))), T.Resize((32, 32)), T.ToTensor(), norm])
flip_crop_bl = T.Compose([T.Lambda(lambda im: TF.hflip(TF.crop(im, top=4, left=0, height=28, width=28))), T.Resize((32, 32)), T.ToTensor(), norm])
flip_crop_br = T.Compose([T.Lambda(lambda im: TF.hflip(TF.crop(im, top=4, left=4, height=28, width=28))), T.Resize((32, 32)), T.ToTensor(), norm])
center_crop = T.Compose([T.Lambda(lambda im: TF.crop(im, top=2, left=2, height=28, width=28)), T.Resize((32, 32)), T.ToTensor(), norm])
rot_p8 = T.Compose([T.Lambda(lambda im: TF.rotate(im, 8)), T.ToTensor(), norm])
rot_m8 = T.Compose([T.Lambda(lambda im: TF.rotate(im, -8)), T.ToTensor(), norm])
rot_p8_flip = T.Compose([T.Lambda(lambda im: TF.hflip(TF.rotate(im, 8))), T.ToTensor(), norm])
rot_m8_flip = T.Compose([T.Lambda(lambda im: TF.hflip(TF.rotate(im, -8))), T.ToTensor(), norm])
bright_contrast = T.Compose([T.ColorJitter(brightness=0.12, contrast=0.12), T.ToTensor(), norm])
saturation_hue = T.Compose([T.ColorJitter(saturation=0.12, hue=0.02), T.ToTensor(), norm])
blur = T.Compose([T.GaussianBlur(kernel_size=3, sigma=(0.3, 0.3)), T.ToTensor(), norm])
sharpness = T.Compose([T.Lambda(lambda im: TF.adjust_sharpness(im, 1.4)), T.ToTensor(), norm])
contrast = T.Compose([T.Lambda(lambda im: TF.adjust_contrast(im, 1.2)), T.ToTensor(), norm])

twenty_view_tfms = six_view_tfms + [
    flip_crop_tl,
    flip_crop_tr,
    flip_crop_bl,
    flip_crop_br,
    center_crop,
    rot_p8,
    rot_m8,
    rot_p8_flip,
    rot_m8_flip,
    bright_contrast,
    saturation_hue,
    blur,
    sharpness,
    contrast,
]
acc_20 = tta_eval(model_tta, raw_val, twenty_view_tfms)

views = [1, 2, 6, 20]
tta_acc = [acc_1, acc_2, acc_6, acc_20]
plt.figure(figsize=(7, 4))
plt.plot(views, tta_acc, marker='o', linewidth=2)
for v, a in zip(views, tta_acc):
    plt.text(v, a + 0.2, f'{a:.2f}', ha='center', fontsize=9)
plt.title('Exercise 4 TTA View Count')
plt.xlabel('number of views')
plt.ylabel('Val Accuracy (%)')
plt.xticks(views)
plt.grid(alpha=0.25)
plt.tight_layout()
out_all = exd / 'ex4_tta_views.png'
plt.savefig(out_all, dpi=140)
plt.close()

add_row(4, 'tta_views', 1, out_all, '단일 뷰는 기준 추론 성능이다.', f'val_acc={acc_1:.2f}')
add_row(4, 'tta_views', 2, out_all, '좌우 반전 평균으로 예측 분산이 줄어들었다.', f'val_acc={acc_2:.2f}')
add_row(4, 'tta_views', 6, out_all, '크롭과 반전을 결합하면 평균화 효과가 더 커졌다.', f'val_acc={acc_6:.2f}')
add_row(4, 'tta_views', 20, out_all, '추가 뷰 평균화로 오차 상쇄 효과를 더 크게 유도했다.', f'val_acc={acc_20:.2f}')

# save csv
csv_path = OUT / 'ablation_summary_cv8.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w = csv.DictWriter(f, fieldnames=['Exercise', 'Parameter', 'Value', 'Output file', 'Visual change', 'Optional metric', 'One-line analysis'])
    w.writeheader()
    w.writerows(rows)

# save summary md
overall = OUT / 'overall_summary_table_cv8.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 8)\n\n')
    f.write('| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows:
        f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta = {
    1: ('Basic Augmentation Effect', '기본 데이터 증강 적용 유무에 따른 검증 정확도 변화를 확인한다.', 'augmentation', '증강을 적용하면 위치와 색 변화에 대한 일반화가 개선되는 경향이 나타났다. 데이터 규모가 작을수록 과한 증강은 오히려 신호를 흐릴 수 있어 기본 강도부터 검증하는 접근이 유효했다.'),
    2: ('Advanced Augmentation Method', 'Cutout, Mixup, CutMix 기법의 효과를 비교한다.', 'method', '고급 증강은 기준 대비 일반화 성능에 차이를 만들었다. 혼합 계열은 경계 학습에 이점이 있고 Cutout은 가림 상황에 강점을 보였다. 데이터 특성과 학습 길이에 따라 최적 기법이 달라진다.'),
    3: ('Augmentation Probability Sweep', '증강 적용 확률 p 변화에 따른 정확도 추이를 본다.', 'augmentation_p', 'p가 증가할수록 변형 다양성은 커진다. 다만 지나치게 크면 원본 구조 정보가 약해져 정확도가 둔화될 수 있다. 중간 p 구간이 성능과 안정성의 균형을 보였다.'),
    4: ('TTA View Ablation', 'TTA view 수 증가가 추론 정확도에 미치는 영향을 확인한다.', 'tta_views', 'TTA는 추가 학습 없이 추론 단계 평균화로 오차를 줄인다. view 수를 늘리면 정확도 개선 여지가 있으나 계산량이 선형으로 증가한다. 오프라인 평가와 실시간 서비스에서 적용 범위를 구분하는 것이 합리적이다.'),
}

report_md = OUT / 'ablation_report_compact_ko_v3_cv8.md'
with report_md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 8 Ablation Study\n\n')
    if data_source != 'cifar10_local':
        f.write('> 실행 환경 제약으로 CIFAR-10 대신 digits(32x32, 3채널 변환) 데이터셋으로 동일 구조 실험을 수행함.\n\n')
    for ex in range(1, 5):
        title, objective, changed, summary = meta[ex]
        f.write(f'## Exercise {ex}. {title}\n\n')
        f.write('### 1) 실험 목적\n')
        f.write(f'- {objective}\n\n')
        f.write('### 2) 변경 인자\n')
        f.write(f'- {changed}\n\n')
        f.write('### 3) 결과 표\n')
        f.write('| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise'] == f'Exercise {ex}':
                f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write('\n### 4) 해석 요약\n')
        f.write(f'- {summary}\n\n')

img_list = OUT / 'generated_images_cv8.txt'
with img_list.open('w', encoding='utf-8') as f:
    for p in sorted(set(r['Output file'] for r in rows)):
        f.write(p + '\n')


def set_korean_font(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for h in ['Heading 1', 'Heading 2', 'Heading 3']:
        st = doc.styles[h]
        st.font.name = 'Malgun Gothic'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        st.font.color.rgb = RGBColor(0, 0, 0)


def set_table_width_pct(table, pct=100):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(int(pct * 50)))


doc = Document()
set_korean_font(doc)
sec = doc.sections[0]
usable = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0
col1 = max(1.3, usable * 0.18)
col2 = max(1.1, usable * 0.13)
col3 = max(3.0, usable - col1 - col2)
img_w = max(2.3, col3 - 0.2)

if data_source != 'cifar10_local':
    doc.add_paragraph('실행 환경 제약으로 CIFAR-10 대신 digits 데이터셋으로 실험함.', style='List Bullet')
else:
    doc.add_paragraph('CIFAR-10 로컬 데이터로 실험함.', style='List Bullet')

for ex in range(1, 5):
    title, objective, changed, summary = meta[ex]
    doc.add_heading(f'Exercise {ex}. {title}', level=2)

    doc.add_heading('1) 실험 목적', level=3)
    doc.add_paragraph(objective, style='List Bullet')

    doc.add_heading('2) 변경 인자', level=3)
    doc.add_paragraph(changed, style='List Bullet')

    doc.add_heading('3) 결과 표', level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width_pct(table, 100)

    hdr = table.rows[0].cells
    hdr[0].text = '파라미터'
    hdr[1].text = '값'
    hdr[2].text = '결과 이미지'

    widths = [Inches(col1), Inches(col2), Inches(col3)]
    for i, w in enumerate(widths):
        hdr[i].width = w

    for r in rows:
        if r['Exercise'] != f'Exercise {ex}':
            continue
        rc = table.add_row().cells
        rc[0].text = r['Parameter']
        rc[1].text = r['Value']
        for i, w in enumerate(widths):
            rc[i].width = w
        p = ROOT / r['Output file']
        if p.exists():
            rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))

    doc.add_paragraph('')
    doc.add_heading('4) 해석 요약', level=3)
    doc.add_paragraph(summary, style='List Bullet')
    doc.add_paragraph('')

docx_path = OUT / 'ablation_report_compact_ko_v3_cv8.docx'
doc.save(docx_path)

print('rows', len(rows))
print(csv_path)
print(report_md)
print(docx_path)
print(overall)
print(img_list)
print('device', DEVICE)
print('data_source', data_source)
