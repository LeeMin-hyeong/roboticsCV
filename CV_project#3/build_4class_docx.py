from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "YOLO11m_4class_operating_report.docx"
REPORT = ROOT / "report_yolo11m"
RUN = ROOT / "runs" / "yolov11_report" / "tools_report"


ACCENT = "1F4E79"
ACCENT_DARK = "17365D"
LIGHT = "EAF2F8"
LIGHTER = "F6F8FA"
GRID = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=GRID):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill=ACCENT):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Cm(0.08)
            cell.margin_bottom = Cm(0.08)
            cell.margin_left = Cm(0.12)
            cell.margin_right = Cm(0.12)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.bold = True
            elif row_idx % 2 == 0:
                set_cell_shading(cell, LIGHTER)


def add_table(doc, headers, rows, widths=None, number_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[idx], ACCENT)
        if widths:
            hdr[idx].width = widths[idx]

    number_cols = set(number_cols or [])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in number_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)
            if widths:
                cells[idx].width = widths[idx]

    style_table(table)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 96, 105)


def add_image(doc, path, caption, width=6.2):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_table_borders(table, color="B7C9D6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def set_document_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.3)

    for name, size, color in [
        ("Heading 1", 17, ACCENT_DARK),
        ("Heading 2", 13.5, ACCENT),
        ("Heading 3", 11.5, ACCENT_DARK),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 9)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("YOLO11m 4-Class Operating Report")
    r.font.size = Pt(25)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = doc.add_paragraph()
    r = p.add_run("Hammer, Pliers, Screwdriver, Wrench 중심 운영 분석")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(92, 103, 115)

    doc.add_paragraph()
    add_callout(
        doc,
        "보고서 목적",
        "학습 산출물을 4개 운영 클래스 기준으로 정리하고, 그래프와 재구성 confusion matrix를 이용해 실제 사용 관점의 성능을 해석한다.",
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["운영 클래스", "Hammer, Pliers, Screwdriver, Wrench"],
            ["권장 가중치", "runs/yolov11_report/tools_report/weights/best.pt"],
            ["기준 성능", "mAP50 0.86456, mAP50-95 0.69299"],
            ["핵심 정책", "Pliers와 plier를 하나의 Pliers 클래스로 통합"],
        ],
        widths=[Cm(4), Cm(11)],
    )
    doc.add_page_break()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    set_document_styles(doc)

    add_cover(doc)

    doc.add_heading("1. 운영 클래스 정의", level=1)
    doc.add_paragraph(
        "본 보고서는 이미 학습이 완료된 YOLO11m 모델 산출물을 바탕으로, 실제 운영 대상을 다음 4개 클래스로 제한해 재구성한 분석 보고서이다."
    )
    add_table(
        doc,
        ["운영 클래스", "클래스 매핑"],
        [
            ["Hammer", "Hammer"],
            ["Pliers", "Pliers, plier"],
            ["Screwdriver", "Screwdriver"],
            ["Wrench", "Wrench"],
        ],
        widths=[Cm(5), Cm(10)],
    )
    doc.add_paragraph(
        "운영 기준은 추론 결과 중 위 4개 클래스만 사용하는 방식으로 정의한다. 특히 Pliers와 plier는 의미상 같은 물체로 보고 하나의 Pliers 운영 클래스로 통합한다."
    )

    doc.add_heading("2. 주요 성능 요약", level=1)
    add_table(
        doc,
        ["지표", "최고값", "Epoch"],
        [
            ["mAP50", "0.86456", "124"],
            ["mAP50-95", "0.69299", "124"],
            ["Precision", "0.88118", "67"],
            ["Recall", "0.83828", "105"],
        ],
        widths=[Cm(5), Cm(4), Cm(3)],
        number_cols=[1, 2],
    )
    add_callout(
        doc,
        "모델 선택",
        "최고 mAP가 124 epoch에서 기록되었으므로 운영 기준 가중치는 last.pt보다 best.pt가 적절하다.",
    )

    doc.add_heading("3. 4개 운영 클래스 데이터 분포", level=1)
    add_table(
        doc,
        ["운영 클래스", "Train annotations", "Val annotations", "Total annotations"],
        [
            ["Hammer", "1,778", "322", "2,100"],
            ["Pliers", "1,706", "455", "2,161"],
            ["Screwdriver", "1,473", "351", "1,824"],
            ["Wrench", "1,592", "362", "1,954"],
            ["합계", "6,549", "1,490", "8,039"],
        ],
        widths=[Cm(4.4), Cm(3.4), Cm(3.4), Cm(3.4)],
        number_cols=[1, 2, 3],
    )
    doc.add_paragraph(
        "4개 운영 클래스는 전체 annotation 9,578개 중 8,039개를 차지한다. 비율로는 약 83.9%이다. 데이터셋의 대부분이 실제 운영 대상 클래스에 해당하므로, 4개 클래스 운영 관점의 보고서로 재구성하는 것은 타당하다."
    )
    add_image(doc, REPORT / "labels.jpg", "Figure 1. Dataset label distribution overview", width=6.0)

    doc.add_heading("4. 학습 곡선 해석", level=1)
    doc.add_heading("4.1 전체 학습 결과", level=2)
    add_image(doc, REPORT / "results.png", "Figure 2. Overall YOLO11m training results", width=6.5)
    doc.add_paragraph(
        "전체 학습 결과는 정상적인 수렴 형태를 보인다. 학습 초반에는 loss가 빠르게 감소하고 precision, recall, mAP가 급격히 상승한다. 이후 100 epoch 전후부터 성능 증가 폭이 줄어들며 안정화된다."
    )

    doc.add_heading("4.2 Train loss", level=2)
    add_image(doc, REPORT / "graph_train_loss.png", "Figure 3. Training loss curve", width=6.5)
    doc.add_paragraph(
        "Train loss는 box_loss, cls_loss, dfl_loss 모두 꾸준히 감소했다. 이는 모델이 학습 데이터에서 객체 위치, 클래스 구분, bounding box 경계 추정을 점진적으로 개선했다는 의미이다."
    )

    doc.add_heading("4.3 Validation loss", level=2)
    add_image(doc, REPORT / "graph_val_loss.png", "Figure 4. Validation loss curve", width=6.5)
    doc.add_paragraph(
        "Validation loss도 전반적으로 감소했다. 이는 학습 데이터에만 맞춰진 과적합보다는 검증 데이터에서도 성능이 개선되었음을 보여준다. 후반부에는 손실이 완만하게 안정화된다."
    )

    doc.add_heading("4.4 mAP", level=2)
    add_image(doc, REPORT / "graph_map.png", "Figure 5. mAP curve", width=6.5)
    doc.add_paragraph(
        "mAP50은 0.86456, mAP50-95는 0.69299까지 상승했다. mAP50은 객체를 대략적으로 잘 찾는지를 보여주고, mAP50-95는 더 엄격한 위치 정확도까지 반영한다."
    )

    doc.add_heading("4.5 Precision / Recall", level=2)
    add_image(doc, REPORT / "graph_precision_recall.png", "Figure 6. Precision and recall curve", width=6.5)
    doc.add_paragraph(
        "Precision과 recall은 모두 0.8 이상 수준으로 안정화되었다. 이는 운영 대상 공구를 찾을 때 오검출과 미검출이 모두 과도하지 않은 균형 상태라는 의미이다."
    )
    add_table(
        doc,
        ["운영 목표", "권장 방향"],
        [
            ["놓치는 객체를 줄이고 싶을 때", "confidence threshold를 낮춰 recall 우선"],
            ["오검출을 줄이고 싶을 때", "confidence threshold를 높여 precision 우선"],
            ["일반적인 자동 인식", "F1 curve 기준 threshold 선택"],
        ],
        widths=[Cm(6.3), Cm(8.0)],
    )

    doc.add_heading("5. 4개 클래스 기준 Confusion Matrix", level=1)
    add_image(
        doc,
        REPORT / "confusion_matrix_4class_reconstructed.png",
        "Figure 7. Reconstructed 4-class confusion matrix",
        width=5.5,
    )
    doc.add_paragraph(
        "위 행렬은 원본 normalized confusion matrix에서 운영 대상 클래스만 추출하고, Pliers와 plier를 하나의 Pliers 클래스로 통합해 재구성한 것이다. 행은 예측 클래스, 열은 실제 클래스를 의미한다."
    )
    add_table(
        doc,
        ["Predicted \\ True", "Hammer", "Pliers", "Screwdriver", "Wrench"],
        [
            ["Hammer", "0.89", "0.00", "0.00", "0.01"],
            ["Pliers", "0.00", "0.96", "0.01", "0.01"],
            ["Screwdriver", "0.01", "0.02", "0.89", "0.01"],
            ["Wrench", "0.00", "0.00", "0.00", "0.89"],
        ],
        widths=[Cm(4.1), Cm(2.6), Cm(2.6), Cm(3.0), Cm(2.6)],
        number_cols=[1, 2, 3, 4],
    )
    doc.add_page_break()
    doc.add_heading("5.1 Confusion Matrix 해석 요약", level=2)
    add_table(
        doc,
        ["운영 클래스", "관찰된 경향"],
        [
            ["Hammer", "대각선 값이 약 0.89로 양호하다. 일부 background 또는 유사 공구와의 혼동이 있으나 주요 클래스 중 안정적인 편이다."],
            ["Pliers", "통합 기준 대각선 값이 약 0.96으로 가장 높다. 두 라벨을 하나의 운영 클래스로 합치면 실제 사용성이 좋아진다."],
            ["Screwdriver", "대각선 값이 약 0.89로 양호하지만, 일부 background 혼동이 보인다. 길고 얇은 물체 특성상 작은 객체나 겹침 상황에서 누락 가능성이 있다."],
            ["Wrench", "대각선 값이 약 0.89로 양호하다. 공구류 중에서도 비교적 안정적으로 분류되는 편이다."],
        ],
        widths=[Cm(3.4), Cm(11.2)],
    )

    doc.add_heading("6. 추론 결과 필터링 정책", level=1)
    add_table(
        doc,
        ["기존 ID", "기존 클래스명", "운영 처리"],
        [
            ["4", "Hammer", "유지"],
            ["7", "Pliers", "Pliers로 통합"],
            ["8", "Screwdriver", "유지"],
            ["10", "Wrench", "유지"],
            ["11", "plier", "Pliers로 통합"],
            ["기타", "Tools, 0, 1, Drill, Hardhat, Measuring Tape, Toolbox", "무시"],
        ],
        widths=[Cm(2.4), Cm(7.0), Cm(5.3)],
    )
    for item in [
        "모델 추론 결과 중 class id가 4, 7, 8, 10, 11인 detection만 남긴다.",
        "class id 7과 11은 모두 Pliers로 표시한다.",
        "그 외 클래스는 화면 표시, 저장, 통계 집계에서 제외한다.",
        "기본 모델은 best.pt를 사용한다.",
        "threshold는 목적에 따라 조정하되, 초기값은 confidence 0.25, IoU 0.7 근처에서 시작한다.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("7. Validation 예측 이미지 활용", level=1)
    add_image(doc, RUN / "val_batch0_labels.jpg", "Figure 8. Validation labels batch 0", width=6.0)
    add_image(doc, RUN / "val_batch0_pred.jpg", "Figure 9. Validation predictions batch 0", width=6.0)
    doc.add_paragraph(
        "Validation 예측 이미지는 4개 운영 클래스만 눈으로 확인할 때 중요하다. 작은 screwdriver나 겹친 wrench처럼 box가 흔들릴 수 있는 사례를 별도로 확인하고, Pliers와 plier는 하나의 Pliers로 합쳐 해석한다."
    )

    doc.add_page_break()
    doc.add_heading("8. 결론", level=1)
    doc.add_paragraph(
        "최종 운영 대상은 Hammer, Pliers, Screwdriver, Wrench 네 클래스이다. 이 네 클래스는 전체 annotation의 약 83.9%를 차지하고, 각 클래스의 학습/검증 데이터도 충분한 편이다."
    )
    doc.add_paragraph(
        "운영 관점의 핵심은 Pliers와 plier를 하나로 합치고, 나머지 클래스 출력은 무시하는 것이다. 이렇게 하면 실제 필요한 공구 4종 탐지에 집중할 수 있다."
    )
    add_table(
        doc,
        ["항목", "권장값"],
        [
            ["weight", "runs/yolov11_report/tools_report/weights/best.pt"],
            ["사용 클래스 ID", "4, 7, 8, 10, 11"],
            ["표시 클래스", "Hammer, Pliers, Screwdriver, Wrench"],
            ["통합 규칙", "Pliers + plier -> Pliers"],
            ["제외 클래스", "Tools, 0, 1, Drill, Hardhat, Measuring Tape, Toolbox"],
        ],
        widths=[Cm(3.7), Cm(10.8)],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
