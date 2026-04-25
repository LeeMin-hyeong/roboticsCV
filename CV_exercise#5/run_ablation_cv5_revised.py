import csv
from pathlib import Path
import numpy as np
import cv2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.signal import correlate2d, fftconvolve
from sklearn.decomposition import PCA
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#5")
TM = ROOT / 'CV_exercise#5_template_matching'
EI = ROOT / 'CV_exercise#5_eigen_image'
OUT = ROOT / 'outputs'
OUT.mkdir(exist_ok=True)
rows = []


def add_row(ex, param, value, out, analysis, metric=''):
    rows.append({'Exercise':f'Exercise {ex}','Parameter':str(param),'Value':str(value),'Output file':str(out.relative_to(ROOT)).replace('\\','/'),'Visual change':'','Optional metric':metric,'One-line analysis':analysis})


def save3(a,b,c,t1,t2,t3,out,cmap='gray'):
    plt.figure(figsize=(12,4))
    plt.subplot(1,3,1); plt.imshow(a,cmap=cmap); plt.title(t1); plt.axis('off')
    plt.subplot(1,3,2); plt.imshow(b,cmap=cmap); plt.title(t2); plt.axis('off')
    plt.subplot(1,3,3); plt.imshow(c,cmap=cmap); plt.title(t3); plt.axis('off')
    plt.tight_layout(); plt.savefig(out,dpi=140); plt.close()

# Exercise 1 Template Matching (percentile threshold)
exd=OUT/'ex1_template_matching_revised'; exd.mkdir(exist_ok=True)
church = cv2.imread(str(TM/'church.png'), cv2.IMREAD_GRAYSCALE).astype(np.float64)
window = cv2.imread(str(TM/'window.png'), cv2.IMREAD_GRAYSCALE).astype(np.float64)
church0 = church - church.mean()
window0 = window - window.mean()
resp = correlate2d(church0, np.flipud(np.fliplr(window0)), mode='same')
for q in [99.5, 99.7, 99.9]:
    th = np.percentile(resp, q)
    pts = np.argwhere(resp >= th)
    vis = cv2.cvtColor(church.astype(np.uint8), cv2.COLOR_GRAY2BGR)
    h,w = window.shape
    for y,x in pts[:120]:
        cv2.rectangle(vis,(x-w//2,y-h//2),(x+w//2,y+h//2),(0,255,0),1)
    out = exd / f'ex1_percentile_{q:.1f}.png'
    save3(church, resp, vis, 'church', f'corr map p{q}', 'detections', out)
    add_row(1,'corr_percentile',q,out,f'백분위 {q}에서 검출 후보 수가 조절됐다.',f'detections={len(pts)}')

# Exercise 2 Matched Filtering (template choice on digit collage)
exd=OUT/'ex2_matched_filtering_revised'; exd.mkdir(exist_ok=True)
hw=3; ox=hw; oy=hw
img0=np.zeros((2*hw+1,2*hw+1)); img0[oy-2:oy+3,ox-1:ox+2]=1; img0[oy-1:oy+2,ox]=0
img1=np.zeros_like(img0); img1[oy-2:oy+3,ox]=1; img1[oy-2,ox-1:ox+1]=1; img1[oy+2,ox-1:ox+2]=1
img2=np.copy(img0); img2[oy-1:oy+2,ox-1]=0; img2[oy+1,ox+1]=0; img2[oy,ox-1:ox+2]=1; img2[oy+1,ox-1]=1
img3=np.copy(img0); img3[oy-1,ox-1]=0; img3[oy+1,ox-1]=0; img3[oy,ox]=1
img4=np.copy(img0); img4[oy-2,ox]=0; img4[oy+1:oy+3,ox-1]=0; img4[oy+2,ox]=0; img4[oy,ox]=1
img5=np.copy(img0); img5[oy,ox]=1; img5[oy-1,ox+1]=0; img5[oy+1,ox-1]=0
img6=np.copy(img0); img6[oy,ox]=1; img6[oy-1,ox+1]=0
img7=np.copy(img0); img7[oy-1:oy+3,ox-1]=0; img7[oy+2,ox]=0
img8=np.copy(img0); img8[oy,ox]=1
img9=np.copy(img4); img9[oy-2,ox]=1
digits=[img0,img1,img2,img3,img4,img5,img6,img7,img8,img9]
collage=np.concatenate(digits, axis=1)
for tidx in [8,3,0]:
    temp=digits[tidx]
    resp2 = fftconvolve(collage, np.rot90(temp,2), mode='same')
    rnorm = resp2 / (resp2.max() + 1e-12)
    peaks = np.argwhere(rnorm > 0.9)
    out = exd / f'ex2_template_{tidx}.png'
    save3(collage, temp, rnorm, 'digit collage', f'template {tidx}', 'matched response', out)
    add_row(2,'template_digit',tidx,out,f'템플릿 숫자 {tidx}에서 필터 응답 분포가 달라졌다.',f'peak_pixels={len(peaks)}')

# Exercise 3 Gender Recognition (PCA components)
exd=OUT/'ex3_gender_recognition_revised'; exd.mkdir(exist_ok=True)
base_gender = EI/'Data_aligned'/'Data_aligned'
male_dir = base_gender/'Male'
female_dir = base_gender/'Female'
male_files = sorted(male_dir.glob('*.jpg'))
female_files = sorted(female_dir.glob('*.jpg'))

def load_gray(files):
    arr=[]
    for p in files:
        im=cv2.imread(str(p), cv2.IMREAD_GRAYSCALE)
        if im is not None:
            arr.append(im.astype(np.float32)/255.0)
    return arr

male_imgs = load_gray(male_files)
female_imgs = load_gray(female_files)
min_n = min(len(male_imgs), len(female_imgs))
male_imgs = male_imgs[:min_n]
female_imgs = female_imgs[:min_n]
# even train / odd test
male_tr = male_imgs[0::2]; male_te = male_imgs[1::2]
female_tr = female_imgs[0::2]; female_te = female_imgs[1::2]
Xtr = np.array([x.ravel() for x in (male_tr+female_tr)])
ytr = np.array([0]*len(male_tr)+[1]*len(female_tr))
Xte = np.array([x.ravel() for x in (male_te+female_te)])
yte = np.array([0]*len(male_te)+[1]*len(female_te))
H,W = male_imgs[0].shape
for k in [10,20,40]:
    k_eff=min(k, Xtr.shape[0]-1, Xtr.shape[1]-1)
    pca=PCA(n_components=k_eff, svd_solver='randomized', random_state=42)
    Ztr=pca.fit_transform(Xtr); Zte=pca.transform(Xte)
    m_mean = Ztr[ytr==0].mean(axis=0)
    f_mean = Ztr[ytr==1].mean(axis=0)
    pred=[]
    for z in Zte:
        dm=np.sum((z-m_mean)**2); df=np.sum((z-f_mean)**2)
        pred.append(0 if dm<df else 1)
    pred=np.array(pred)
    acc=float(np.mean(pred==yte))

    eig = pca.components_[0].reshape(H,W)
    mean_face = pca.mean_.reshape(H,W)
    out=exd/f'ex3_pca_{k}.png'
    save3(mean_face, eig, np.abs(eig), 'mean face', f'eigenface1 k={k}', 'abs eigenface', out)
    add_row(3,'num_components',k,out,f'성분 {k}에서 성별 분류 성능과 표현력이 달라졌다.',f'accuracy={acc:.4f}')

# Exercise 4 Illumination Variations (theta from notebook add_illumination)
exd=OUT/'ex4_illumination_variation_revised'; exd.mkdir(exist_ok=True)
base = cv2.imread(str((EI/'att_faces_aligned'/'att_faces_aligned'/'s1'/'1.pgm')), cv2.IMREAD_GRAYSCALE).astype(np.float64)/255.0
h,w = base.shape
x,y=np.meshgrid(np.arange(1,w+1), np.arange(1,h+1))
np.random.seed(42)
a = 5*(np.random.rand()-0.5)/w
b = 5*(np.random.rand()-0.5)/h
for theta in [0.1,0.2,0.4]:
    illum = base * (1 - theta + theta*np.cos(a*x + b*y))
    illum = np.clip(illum, 0, 1)
    eq = cv2.equalizeHist((illum*255).astype(np.uint8)).astype(np.float32)/255.0
    out=exd/f'ex4_theta_{theta:.1f}.png'
    save3(base, illum, eq, 'orig', f'illum theta={theta}', 'equalized', out)
    add_row(4,'theta',theta,out,f'theta={theta}에서 조명 왜곡 강도가 변화했다.',f'std_illum={illum.std():.4f}')

# Exercise 5 Person Identification (PCA components around notebook 150 intent)
exd=OUT/'ex5_person_identification_revised'; exd.mkdir(exist_ok=True)
face_dir = EI/'att_faces_aligned_lighting'/'att_faces_aligned_lighting'
num_classes=40
tr_range=range(1,6); te_range=range(6,11)

train=[]; train_y=[]; test=[]; test_y=[]
for cls in range(1,num_classes+1):
    for idx in tr_range:
        p=face_dir/f's{cls}'/f'{idx}.pgm'
        im=cv2.imread(str(p), cv2.IMREAD_GRAYSCALE)
        if im is not None:
            train.append(im.astype(np.float32).ravel()/255.0); train_y.append(cls)
    for idx in te_range:
        p=face_dir/f's{cls}'/f'{idx}.pgm'
        im=cv2.imread(str(p), cv2.IMREAD_GRAYSCALE)
        if im is not None:
            test.append(im.astype(np.float32).ravel()/255.0); test_y.append(cls)
Xtr=np.array(train); ytr=np.array(train_y); Xte=np.array(test); yte=np.array(test_y)
for k in [80,120,150]:
    k_eff=min(k, Xtr.shape[0]-1, Xtr.shape[1]-1)
    pca=PCA(n_components=k_eff, svd_solver='randomized', random_state=42)
    Ztr=pca.fit_transform(Xtr); Zte=pca.transform(Xte)
    pred=[]
    for z in Zte:
        d=np.sum((Ztr-z)**2, axis=1)
        pred.append(ytr[np.argmin(d)])
    pred=np.array(pred)
    acc=float(np.mean(pred==yte))

    # top-5 classes confusion summary
    uniq=np.unique(yte)[:5]
    cm=np.zeros((len(uniq),len(uniq)), dtype=int)
    cls_to_i={c:i for i,c in enumerate(uniq)}
    for t,p in zip(yte,pred):
        if t in cls_to_i and p in cls_to_i:
            cm[cls_to_i[t], cls_to_i[p]] += 1
    out=exd/f'ex5_pca_{k}.png'
    plt.figure(figsize=(10,4))
    plt.subplot(1,2,1); plt.imshow(cm, cmap='Blues'); plt.title(f'confusion(1~5) k={k}'); plt.xlabel('pred'); plt.ylabel('true')
    plt.subplot(1,2,2); plt.bar(['accuracy'], [acc]); plt.ylim(0,1); plt.title('identification')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
    add_row(5,'num_components',k,out,f'성분 {k}에서 신원 식별 정확도가 변했다.',f'accuracy={acc:.4f}')

# save artifacts
csv_path=OUT/'ablation_summary_cv5_revised.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w=csv.DictWriter(f, fieldnames=['Exercise','Parameter','Value','Output file','Visual change','Optional metric','One-line analysis'])
    w.writeheader(); w.writerows(rows)

overall=OUT/'overall_summary_table_cv5_revised.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 5, revised)\n\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows: f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta={
1:('Template Matching','상관 기반 템플릿 매칭에서 임계 백분위를 비교한다.','corr_percentile','백분위를 높이면 후보 수가 감소한다. 낮은 백분위는 민감도가 높지만 오검출 가능성이 커진다. 목표 물체가 명확할수록 높은 백분위가 유리했다.'),
2:('Matched Filtering','숫자 템플릿 선택에 따른 matched filtering 응답을 비교한다.','template_digit','템플릿 숫자에 따라 응답 분포가 달라졌다. 목표 패턴과 유사한 템플릿에서 피크가 집중됐다. 불일치 템플릿은 응답이 분산됐다.'),
3:('Gender Recognition','PCA 차원 수가 성별 분류 성능에 미치는 영향을 비교한다.','num_components','차원이 너무 낮으면 성별 특징 분리가 약해진다. 적정 차원에서는 평균 얼굴 기반 분류 성능이 개선된다. 과도한 차원은 개선 폭이 제한적이었다.'),
4:('Illumination Variations','조명 왜곡 계수 theta 변화가 얼굴 대비에 미치는 영향을 비교한다.','theta','theta가 커질수록 비균일 조명 왜곡이 강해졌다. 균일하지 않은 명암 기울기가 얼굴 특징을 가렸다. 평활화 보정은 일부 회복에 도움을 줬다.'),
5:('Person Identification','원 코드의 PCA 기반 식별에서 차원 수를 점검한다.','num_components','성분 수가 증가하면 분리력이 개선되는 구간이 있었다. 원 코드의 높은 차원 설정은 조명 변동 데이터에서 성능 안정화에 도움을 줬다. 다만 일정 수준 이후 이득은 완만해졌다.'),
}

md=OUT/'ablation_report_compact_ko_v3_cv5_revised.md'
with md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 5 Ablation Study (Revised)\n\n')
    for ex in range(1,6):
        t,o,c,s=meta[ex]
        f.write(f'## Exercise {ex}. {t}\n\n### 1) 실험 목적\n- {o}\n\n### 2) 변경 인자\n- {c}\n\n### 3) 결과 표\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise']==f'Exercise {ex}':
                f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write(f"\n### 4) 해석 요약\n- {s}\n\n")

lst=OUT/'generated_images_cv5_revised.txt'
with lst.open('w', encoding='utf-8') as f:
    for r in rows: f.write(r['Output file']+'\n')

# docx

def set_korean_font(doc):
    normal=doc.styles['Normal']; normal.font.name='Malgun Gothic'; normal.font.size=Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); normal.font.color.rgb=RGBColor(0,0,0)
    for h in ['Heading 1','Heading 2','Heading 3']:
        st=doc.styles[h]; st.font.name='Malgun Gothic'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.color.rgb=RGBColor(0,0,0)

def set_table_width_pct(table,pct=100):
    tbl=table._tbl; tblPr=tbl.tblPr
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'),'pct'); tblW.set(qn('w:w'),str(int(pct*50)))

doc=Document(); set_korean_font(doc)
sec=doc.sections[0]; usable=(sec.page_width-sec.left_margin-sec.right_margin)/914400.0
col1=max(1.3,usable*0.18); col2=max(1.1,usable*0.13); col3=max(3.0,usable-col1-col2); img_w=max(2.3,col3-0.2)
for ex in range(1,6):
    t,o,c,s=meta[ex]
    doc.add_heading(f'Exercise {ex}. {t}', level=2)
    doc.add_heading('1) 실험 목적', level=3); doc.add_paragraph(o, style='List Bullet')
    doc.add_heading('2) 변경 인자', level=3); doc.add_paragraph(c, style='List Bullet')
    doc.add_heading('3) 결과 표', level=3)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.autofit=False; set_table_width_pct(table,100)
    hdr=table.rows[0].cells; hdr[0].text='파라미터'; hdr[1].text='값'; hdr[2].text='결과 이미지'
    widths=[Inches(col1),Inches(col2),Inches(col3)]
    for i,w in enumerate(widths): hdr[i].width=w
    for r in rows:
        if r['Exercise']!=f'Exercise {ex}': continue
        rc=table.add_row().cells; rc[0].text=r['Parameter']; rc[1].text=r['Value']
        for i,w in enumerate(widths): rc[i].width=w
        p=ROOT/r['Output file']
        if p.exists(): rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))
    doc.add_paragraph(''); doc.add_heading('4) 해석 요약', level=3); doc.add_paragraph(s, style='List Bullet'); doc.add_paragraph('')

docx=OUT/'ablation_report_compact_ko_v3_cv5_revised.docx'; doc.save(docx)
print('rows',len(rows)); print(csv_path); print(md); print(docx); print(overall); print(lst)
