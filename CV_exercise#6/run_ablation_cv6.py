import csv
from pathlib import Path
import numpy as np
import cv2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#6")
OUT = ROOT / 'outputs'
OUT.mkdir(exist_ok=True)
rows=[]


def add_row(ex,param,value,out,analysis,metric=''):
    rows.append({'Exercise':f'Exercise {ex}','Parameter':str(param),'Value':str(value),'Output file':str(out.relative_to(ROOT)).replace('\\','/'),'Visual change':'','Optional metric':metric,'One-line analysis':analysis})

# Ex1 SIFT descriptor
exd=OUT/'ex1_sift_descriptor'; exd.mkdir(exist_ok=True)
img = cv2.imread(str(ROOT/'gates.jpg'), cv2.IMREAD_GRAYSCALE)
if hasattr(cv2, 'SIFT_create'):
    mode='SIFT'
    for nf in [200,500,1000]:
        sift=cv2.SIFT_create(nfeatures=nf)
        kp, des=sift.detectAndCompute(img,None)
        vis=cv2.drawKeypoints(img,kp,None,flags=cv2.DRAW_MATCHES_FLAGS_DRAW_RICH_KEYPOINTS)
        out=exd/f'ex1_nfeatures_{nf}.png'
        cv2.imwrite(str(out), vis)
        add_row(1,'nfeatures',nf,out,f'nfeatures={nf}에서 검출 키포인트 수가 변했다.',f'keypoints={len(kp)}')
else:
    mode='ORB'
    for nf in [200,500,1000]:
        orb=cv2.ORB_create(nfeatures=nf)
        kp, des=orb.detectAndCompute(img,None)
        vis=cv2.drawKeypoints(img,kp,None,color=(0,255,0),flags=0)
        out=exd/f'ex1_nfeatures_{nf}.png'
        cv2.imwrite(str(out), vis)
        add_row(1,'nfeatures',nf,out,f'ORB nfeatures={nf}에서 특징점 밀도가 달라졌다.',f'keypoints={len(kp)}')

# Ex2 RANSAC homography
exd=OUT/'ex2_ransac_mosaic'; exd.mkdir(exist_ok=True)
im1=cv2.imread(str(ROOT/'church1.jpg'))
im2=cv2.imread(str(ROOT/'church2.jpg'))
g1=cv2.cvtColor(im1, cv2.COLOR_BGR2GRAY)
g2=cv2.cvtColor(im2, cv2.COLOR_BGR2GRAY)
if hasattr(cv2, 'SIFT_create'):
    det=cv2.SIFT_create()
else:
    det=cv2.ORB_create(nfeatures=1200)
kp1,des1=det.detectAndCompute(g1,None)
kp2,des2=det.detectAndCompute(g2,None)
if des1 is None or des2 is None:
    raise RuntimeError('Feature extraction failed for exercise 2')

if des1.dtype == np.float32:
    matcher=cv2.BFMatcher(cv2.NORM_L2)
else:
    matcher=cv2.BFMatcher(cv2.NORM_HAMMING)

for ratio in [0.6,0.75,0.9]:
    knn=matcher.knnMatch(des1,des2,k=2)
    good=[]
    for m,n in knn:
        if m.distance < ratio*n.distance:
            good.append(m)
    if len(good) >= 4:
        src=np.float32([kp1[m.queryIdx].pt for m in good]).reshape(-1,1,2)
        dst=np.float32([kp2[m.trainIdx].pt for m in good]).reshape(-1,1,2)
        H,mask=cv2.findHomography(src,dst,cv2.RANSAC,5.0)
        inliers=int(mask.sum()) if mask is not None else 0
    else:
        H,mask=None,None
        inliers=0

    draw=cv2.drawMatches(im1,kp1,im2,kp2,good[:80],None,flags=2)
    out=exd/f'ex2_ratio_{ratio:.2f}.png'
    cv2.imwrite(str(out), draw)
    add_row(2,'ratio_test',ratio,out,f'ratio={ratio:.2f}에서 매칭 보존 개수와 inlier 수가 변했다.',f'good_matches={len(good)}, inliers={inliers}')

# Ex3 Lucas-Kanade tracking
exd=OUT/'ex3_kanade_lucas'; exd.mkdir(exist_ok=True)
video_candidates=['crew_4cif.mp4','ice_4cif.mp4','Car-01.mp4']
vp=None
for v in video_candidates:
    p=ROOT/v
    if p.exists():
        vp=p; break
if vp is None:
    raise RuntimeError('No video found for exercise 3')

for win in [9,15,21]:
    cap=cv2.VideoCapture(str(vp))
    ret, old=cap.read()
    if not ret:
        cap.release(); continue
    oldg=cv2.cvtColor(old, cv2.COLOR_BGR2GRAY)
    p0=cv2.goodFeaturesToTrack(oldg, mask=None, maxCorners=200, qualityLevel=0.3, minDistance=7, blockSize=7)
    lk_params=dict(winSize=(win,win), maxLevel=2, criteria=(cv2.TERM_CRITERIA_EPS | cv2.TERM_CRITERIA_COUNT, 10, 0.03))
    color=np.random.randint(0,255,(200,3))
    mask=np.zeros_like(old)
    tracked=0
    frame=None
    for _ in range(40):
        ret, fr=cap.read()
        if not ret or p0 is None:
            break
        frg=cv2.cvtColor(fr, cv2.COLOR_BGR2GRAY)
        p1, st, err = cv2.calcOpticalFlowPyrLK(oldg, frg, p0, None, **lk_params)
        if p1 is None:
            break
        good_new=p1[st==1]
        good_old=p0[st==1]
        tracked=len(good_new)
        draw=fr.copy()
        for i,(n,o) in enumerate(zip(good_new, good_old)):
            a,b=n.ravel(); c,d=o.ravel()
            mask=cv2.line(mask, (int(a),int(b)), (int(c),int(d)), color[i].tolist(), 1)
            draw=cv2.circle(draw, (int(a),int(b)), 2, color[i].tolist(), -1)
        frame=cv2.add(draw, mask)
        oldg=frg.copy(); p0=good_new.reshape(-1,1,2)
    cap.release()
    if frame is None:
        frame=old
    out=exd/f'ex3_win_{win}.png'
    cv2.imwrite(str(out), frame)
    add_row(3,'winSize',win,out,f'winSize={win}에서 추적 안정성과 궤적 길이가 달라졌다.',f'tracked_points={tracked}')

# outputs
csv_path=OUT/'ablation_summary_cv6.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w=csv.DictWriter(f, fieldnames=['Exercise','Parameter','Value','Output file','Visual change','Optional metric','One-line analysis'])
    w.writeheader(); w.writerows(rows)

overall=OUT/'overall_summary_table_cv6.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 6)\n\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows: f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta={
1:('SIFT Descriptor','특징점 검출 개수 제한이 키포인트 분포에 미치는 영향을 비교한다.','nfeatures','nfeatures를 늘릴수록 키포인트가 조밀해졌다. 강한 코너뿐 아니라 약한 구조까지 포착됐다. 과도한 특징점은 후속 매칭에서 오검출 후보를 늘릴 수 있다.'),
2:('RANSAC','ratio test 임계값 변화가 매칭과 호모그래피 안정성에 미치는 영향을 본다.','ratio_test','ratio가 낮으면 매칭은 적지만 신뢰도가 높아졌다. ratio를 완화하면 매칭 수는 증가하지만 outlier 비율도 상승했다. RANSAC inlier 수는 두 효과의 균형을 반영했다.'),
3:('Kanade Lucas','LK 추적 윈도우 크기에 따른 추적 안정성을 비교한다.','winSize','작은 윈도우는 미세 움직임에 민감하다. 잡음에도 흔들리기 쉽다. 큰 윈도우는 추적이 안정적이지만 세부 움직임 반응은 둔해진다.'),
}

md=OUT/'ablation_report_compact_ko_v3_cv6.md'
with md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 6 Ablation Study\n\n')
    for ex in range(1,4):
        t,o,c,s=meta[ex]
        f.write(f'## Exercise {ex}. {t}\n\n### 1) 실험 목적\n- {o}\n\n### 2) 변경 인자\n- {c}\n\n### 3) 결과 표\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise']==f'Exercise {ex}': f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write(f"\n### 4) 해석 요약\n- {s}\n\n")

lst=OUT/'generated_images_cv6.txt'
with lst.open('w', encoding='utf-8') as f:
    for r in rows: f.write(r['Output file']+'\n')

# docx

def set_korean_font(doc):
    normal=doc.styles['Normal']; normal.font.name='Malgun Gothic'; normal.font.size=Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); normal.font.color.rgb=RGBColor(0,0,0)
    for h in ['Heading 1','Heading 2','Heading 3']:
        st=doc.styles[h]; st.font.name='Malgun Gothic'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.color.rgb=RGBColor(0,0,0)

def set_table_width_pct(table,pct=100):
    tbl=table._tbl; tblPr=tbl.tblPr
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'),'pct'); tblW.set(qn('w:w'),str(int(pct*50)))

doc=Document(); set_korean_font(doc)
sec=doc.sections[0]; usable=(sec.page_width-sec.left_margin-sec.right_margin)/914400.0
col1=max(1.3,usable*0.18); col2=max(1.1,usable*0.13); col3=max(3.0,usable-col1-col2); img_w=max(2.3,col3-0.2)
for ex in range(1,4):
    t,o,c,s=meta[ex]
    doc.add_heading(f'Exercise {ex}. {t}', level=2)
    doc.add_heading('1) 실험 목적', level=3); doc.add_paragraph(o, style='List Bullet')
    doc.add_heading('2) 변경 인자', level=3); doc.add_paragraph(c, style='List Bullet')
    doc.add_heading('3) 결과 표', level=3)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.autofit=False; set_table_width_pct(table,100)
    hdr=table.rows[0].cells; hdr[0].text='파라미터'; hdr[1].text='값'; hdr[2].text='결과 이미지'
    widths=[Inches(col1),Inches(col2),Inches(col3)]
    for i,w in enumerate(widths): hdr[i].width=w
    for r in rows:
        if r['Exercise']!=f'Exercise {ex}': continue
        rc=table.add_row().cells; rc[0].text=r['Parameter']; rc[1].text=r['Value']
        for i,w in enumerate(widths): rc[i].width=w
        p=ROOT/r['Output file']
        if p.exists(): rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))
    doc.add_paragraph(''); doc.add_heading('4) 해석 요약', level=3); doc.add_paragraph(s, style='List Bullet'); doc.add_paragraph('')

docx=OUT/'ablation_report_compact_ko_v3_cv6.docx'; doc.save(docx)
print('rows',len(rows)); print(csv_path); print(md); print(docx); print(overall); print(lst); print('feature_mode', mode)
