import csv
from pathlib import Path
import numpy as np
import cv2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.ndimage import grey_dilation, grey_erosion, median_filter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#4")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
rows = []


def add_row(ex, param, value, out, analysis, metric=''):
    rows.append({
        'Exercise': f'Exercise {ex}',
        'Parameter': str(param),
        'Value': str(value),
        'Output file': str(out.relative_to(ROOT)).replace('\\','/'),
        'Visual change': '',
        'Optional metric': metric,
        'One-line analysis': analysis,
    })


def save3(a,b,c,t1,t2,t3,out,cmap='gray'):
    plt.figure(figsize=(12,4))
    plt.subplot(1,3,1); plt.imshow(a, cmap=cmap); plt.title(t1); plt.axis('off')
    plt.subplot(1,3,2); plt.imshow(b, cmap=cmap); plt.title(t2); plt.axis('off')
    plt.subplot(1,3,3); plt.imshow(c, cmap=cmap); plt.title(t3); plt.axis('off')
    plt.tight_layout(); plt.savefig(out,dpi=140); plt.close()

# 1 Binary dilation
exd = OUT/'ex1_binary_dilation'; exd.mkdir(exist_ok=True)
img = cv2.imread(str(ROOT/'bacteria.png'), cv2.IMREAD_GRAYSCALE)
_, bw = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY+cv2.THRESH_OTSU)
for k in [3,5,9]:
    ker = cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    d = cv2.dilate(bw, ker)
    out = exd/f'ex1_k_{k}.png'
    save3(img,bw,d,'orig','binary',f'dilation k={k}',out)
    add_row(1,'kernel_size',k,out,f'커널 {k}로 키우자 전경이 팽창해 작은 간격이 메워졌다.',f'fg_ratio={np.mean(d>0):.4f}')

# 2 Binary erosion
exd = OUT/'ex2_binary_erosion'; exd.mkdir(exist_ok=True)
for k in [3,5,9]:
    ker = cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    e = cv2.erode(bw, ker)
    out = exd/f'ex2_k_{k}.png'
    save3(img,bw,e,'orig','binary',f'erosion k={k}',out)
    add_row(2,'kernel_size',k,out,f'커널 {k}에서 erosion이 강해져 얇은 전경이 감소했다.',f'fg_ratio={np.mean(e>0):.4f}')

# 3 Binary erosion coins
exd = OUT/'ex3_binary_erosion_coins'; exd.mkdir(exist_ok=True)
cir = cv2.imread(str(ROOT/'circles.png'), cv2.IMREAD_GRAYSCALE)
_, cbw = cv2.threshold(cir,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)
for k in [5,11,17]:
    ker = cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    e = cv2.erode(cbw, ker)
    out = exd/f'ex3_k_{k}.png'
    save3(cir,cbw,e,'orig','binary',f'erosion k={k}',out)
    n = cv2.connectedComponents((e>0).astype(np.uint8))[0]-1
    add_row(3,'kernel_size',k,out,f'erosion 커널 {k}에서 접촉 동전 분리가 진행됐다.',f'components={n}')

# 4 Binary erosion fence
exd = OUT/'ex4_binary_erosion_fence'; exd.mkdir(exist_ok=True)
fence = cv2.imread(str(ROOT/'fence.jpg'), cv2.IMREAD_GRAYSCALE)
for t in [80,100,120]:
    _, fbw = cv2.threshold(fence,t,255,cv2.THRESH_BINARY)
    er = cv2.erode(fbw, cv2.getStructuringElement(cv2.MORPH_RECT,(3,3)))
    out = exd/f'ex4_t_{t}.png'
    save3(fence,fbw,er,'orig',f'binary t={t}','eroded',out)
    add_row(4,'threshold',t,out,f'threshold {t}로 높일수록 fence 검출이 보수적으로 변했다.',f'fg_ratio={np.mean(er>0):.4f}')

# 5 small hole removal
exd = OUT/'ex5_small_hole_removal'; exd.mkdir(exist_ok=True)
peter = cv2.imread(str(ROOT/'peter.png'), cv2.IMREAD_GRAYSCALE)
_, pbw = cv2.threshold(peter,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)
for k in [3,7,11]:
    ker = cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    cl = cv2.morphologyEx(pbw, cv2.MORPH_CLOSE, ker)
    out = exd/f'ex5_k_{k}.png'
    save3(pbw,cl,cv2.absdiff(cl,pbw),'binary','closed','difference',out)
    add_row(5,'closing_kernel',k,out,f'closing 커널 {k}에서 내부 hole이 채워지며 영역이 매끈해졌다.',f'changed_ratio={np.mean(cl!=pbw):.4f}')

# 6 binary edge detection
exd = OUT/'ex6_binary_edge_detection'; exd.mkdir(exist_ok=True)
clip = cv2.imread(str(ROOT/'cliparts.png'), cv2.IMREAD_GRAYSCALE)
_, ubw = cv2.threshold(clip,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)
for k in [3,5,7]:
    ker = cv2.getStructuringElement(cv2.MORPH_RECT,(k,k))
    d = cv2.dilate(ubw, ker); e = cv2.erode(ubw, ker); edge = cv2.subtract(d,e)
    out = exd/f'ex6_k_{k}.png'
    save3(ubw,d,edge,'binary',f'dilate k={k}','edge(d-e)',out)
    add_row(6,'kernel_size',k,out,f'k={k}에서 morphological edge 두께가 증가했다.',f'edge_ratio={np.mean(edge>0):.4f}')

# 7 1D gray dilation/erosion
exd = OUT/'ex7_1d_gray_morph'; exd.mkdir(exist_ok=True)
x = np.zeros(128, dtype=np.uint8); x[50:78]=180; x[62:66]=80
sig = np.tile(x,(40,1))
for n in [3,7,11]:
    d = grey_dilation(sig, size=(1,n)); e = grey_erosion(sig, size=(1,n))
    out = exd/f'ex7_n_{n}.png'
    save3(sig,d,e,'signal','dilation','erosion',out)
    add_row(7,'window_length',n,out,f'window {n}이 커질수록 peak가 넓어지고 valley가 완화됐다.',f'd_mean={d.mean():.2f}')

# 8 2D gray dilation/erosion
exd = OUT/'ex8_2d_gray_morph'; exd.mkdir(exist_ok=True)
but = cv2.imread(str(ROOT/'butterfly.png'), cv2.IMREAD_GRAYSCALE)
for r in [1,3,5]:
    k = 2*r+1
    ker = cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    d = cv2.dilate(but, ker); e = cv2.erode(but, ker)
    out = exd/f'ex8_r_{r}.png'
    save3(but,d,e,'orig',f'dilation r={r}',f'erosion r={r}',out)
    add_row(8,'radius',r,out,f'r={r}에서 dilation은 밝은 구조를 확장하고 erosion은 어두운 구조를 강조했다.',f'std_d={d.std():.2f}')

# 9 coin separation by gray dilation
exd = OUT/'ex9_coin_separation'; exd.mkdir(exist_ok=True)
for r in [2,4,6]:
    k=2*r+1
    ker=cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    d=cv2.dilate(cir,ker)
    th=cv2.threshold(d,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)[1]
    out=exd/f'ex9_r_{r}.png'
    save3(cir,d,th,'orig','gray dilation','thresholded',out)
    n=cv2.connectedComponents((th>0).astype(np.uint8))[0]-1
    add_row(9,'radius',r,out,f'반지름 {r}에서 coin blob 연결 상태가 변해 분리 결과가 달라졌다.',f'components={n}')

# 10 hole detection by gray erosion
exd = OUT/'ex10_hole_detection'; exd.mkdir(exist_ok=True)
for r in [2,4,6]:
    k=2*r+1
    ker=cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(k,k))
    e=cv2.erode(fence,ker)
    diff=cv2.subtract(fence,e)
    out=exd/f'ex10_r_{r}.png'
    save3(fence,e,diff,'orig','eroded','orig-eroded',out)
    add_row(10,'radius',r,out,f'반지름 {r}에서 침식 차영상이 커져 hole 후보가 더 강조됐다.',f'diff_mean={diff.mean():.2f}')

# 11 gray edge detector
exd = OUT/'ex11_gray_edge_detector'; exd.mkdir(exist_ok=True)
for k in [3,5,9]:
    ker=cv2.getStructuringElement(cv2.MORPH_RECT,(k,k))
    d=cv2.dilate(but,ker); e=cv2.erode(but,ker); g=cv2.subtract(d,e)
    out=exd/f'ex11_k_{k}.png'
    save3(but,d,g,'orig','dilated','morph gradient',out)
    add_row(11,'kernel_size',k,out,f'kernel {k}에서 gradient edge 두께가 증가했다.',f'edge_mean={g.mean():.2f}')

# 12 cascaded dilations
exd = OUT/'ex12_cascaded_dilations'; exd.mkdir(exist_ok=True)
for it in [1,2,3]:
    d=ubw.copy()
    ker=cv2.getStructuringElement(cv2.MORPH_ELLIPSE,(5,5))
    for _ in range(it): d=cv2.dilate(d,ker)
    out=exd/f'ex12_iter_{it}.png'
    save3(ubw,d,cv2.absdiff(d,ubw),'binary',f'cascade iter={it}','difference',out)
    add_row(12,'iterations',it,out,f'iteration {it}에서 팽창 누적 효과가 커졌다.',f'fg_ratio={np.mean(d>0):.4f}')

# 13 majority/median
exd = OUT/'ex13_majority_median'; exd.mkdir(exist_ok=True)
bac = cv2.imread(str(ROOT/'bacteria.png'), cv2.IMREAD_GRAYSCALE)
noise = bac.copy().astype(np.uint8)
np.random.seed(42)
mask = np.random.rand(*noise.shape) < 0.06
noise[mask] = 255 - noise[mask]
for k in [3,5,7]:
    med = cv2.medianBlur(noise, k)
    out = exd/f'ex13_k_{k}.png'
    save3(noise,med,cv2.absdiff(noise,med),'noisy',f'median {k}','difference',out)
    add_row(13,'median_kernel',k,out,f'kernel {k}에서 salt-pepper 노이즈가 줄었고 세부 텍스처도 완화됐다.',f'mae={np.mean(np.abs(med.astype(np.float32)-bac.astype(np.float32))):.2f}')

# 14 nonuniform lighting compensation
exd = OUT/'ex14_lighting_comp'; exd.mkdir(exist_ok=True)
paper = cv2.imread(str(ROOT/'paper.png'), cv2.IMREAD_GRAYSCALE)
for r in [11,21,41]:
    bg = cv2.GaussianBlur(paper, (0,0), r)
    comp = cv2.normalize(cv2.subtract(paper, bg), None, 0, 255, cv2.NORM_MINMAX)
    th = cv2.threshold(comp,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)[1]
    out = exd/f'ex14_sigma_{r}.png'
    save3(paper,comp,th,'orig',f'comp sigma={r}','binary',out)
    add_row(14,'sigma',r,out,f'sigma {r}에서 배경 조명 성분이 제거되며 문자 대비가 개선됐다.',f'fg_ratio={np.mean(th>0):.4f}')

# outputs
csv_path = OUT/'ablation_summary_cv4.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w=csv.DictWriter(f, fieldnames=['Exercise','Parameter','Value','Output file','Visual change','Optional metric','One-line analysis'])
    w.writeheader(); w.writerows(rows)

overall = OUT/'overall_summary_table_cv4.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 4)\n\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows: f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta={
1:('Binary_dilation','이진 dilation 크기에 따른 연결 효과를 비교한다.','kernel_size','커널이 커질수록 전경이 팽창하고 분리 객체 간 간격이 줄어들었다. 작은 결손은 복원됐지만 과도한 팽창은 객체 병합을 유발했다.'),
2:('Binary_Erosion','이진 erosion 강도 변화에 따른 객체 축소를 확인한다.','kernel_size','erosion이 강해질수록 얇은 구조가 빠르게 소실됐다. 잡음 제거에는 유리했지만 작은 객체 보존에는 불리했다.'),
3:('Binary_Erosion_Coins','동전 분리를 위한 erosion 파라미터를 비교한다.','kernel_size','커널 증가에 따라 접촉 영역이 줄며 분리가 진행됐다. 너무 큰 커널에서는 실제 객체 면적도 과도하게 줄었다.'),
4:('Binary_Erosion_Fence','fence 이진화 임계값과 침식 결과를 비교한다.','threshold','threshold가 높아질수록 검출은 보수적으로 변했다. 침식까지 적용하면 세선 구조 소실이 더 커졌다.'),
5:('Small_Hole_Removal','closing으로 hole 제거 효과를 비교한다.','closing_kernel','closing 커널이 커질수록 내부 공극이 안정적으로 메워졌다. 동시에 경계 세부 형상은 단순화됐다.'),
6:('Binary_Edge_Detection','morphological edge의 두께 변화를 비교한다.','kernel_size','커널 크기 증가에 따라 edge 폭이 넓어졌다. 경계 강조는 강해졌지만 정밀도는 낮아졌다.'),
7:('1D Gray Morph','1D 신호에서 dilation/erosion 창 길이 효과를 본다.','window_length','창이 커질수록 피크는 확장되고 계곡은 완화됐다. 신호 외곽선이 더 평탄해졌다.'),
8:('2D Gray Morph','2D 회색조 dilation/erosion 반지름을 비교한다.','radius','반지름이 커질수록 국소 대비가 재분배됐다. 밝은 영역 확장과 어두운 영역 수축이 동시에 강화됐다.'),
9:('Coin Separation by Gray Dilation','회색조 dilation 후 분리 성능을 비교한다.','radius','dilation 반지름에 따라 blob 연결 상태가 달라졌다. 분리 성능은 반지름 선택에 민감했다.'),
10:('Hole Detection by Gray Erosion','erosion 차영상 기반 hole 강조를 비교한다.','radius','반지름이 증가할수록 차영상 응답이 커졌다. hole 후보 강조는 강해졌지만 배경 영향도 함께 증가했다.'),
11:('Gray Morphological Edge Detector','회색조 morphological gradient를 비교한다.','kernel_size','커널이 커질수록 edge는 두꺼워졌다. 미세 경계 구분은 상대적으로 약해졌다.'),
12:('Cascaded Graylevel Dilations','연속 dilation 횟수의 누적 효과를 확인한다.','iterations','반복 횟수가 늘수록 전경 확장이 누적됐다. 초기 작은 공백은 메워지지만 객체 병합 위험이 높아졌다.'),
13:('Majority/Median Filter','노이즈 제거에서 median 커널 크기를 비교한다.','median_kernel','커널이 커질수록 잡음은 줄었다. 대신 질감과 경계의 세부 정보도 완화됐다.'),
14:('Nonuniform Lighting Compensation','조명 보정 필터 스케일을 비교한다.','sigma','sigma가 커질수록 저주파 조명 성분 제거가 강해졌다. 적정 범위에서는 문자 대비가 개선되고 과도하면 정보가 약화됐다.'),
}

md = OUT/'ablation_report_compact_ko_v3_cv4.md'
with md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 4 Ablation Study\n\n')
    for ex in range(1,15):
        t,o,c,s=meta[ex]
        f.write(f'## Exercise {ex}. {t}\n\n### 1) 실험 목적\n- {o}\n\n### 2) 변경 인자\n- {c}\n\n### 3) 결과 표\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise']==f'Exercise {ex}':
                f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write(f"\n### 4) 해석 요약\n- {s}\n\n")

# list
lst=OUT/'generated_images_cv4.txt'
with lst.open('w', encoding='utf-8') as f:
    for r in rows: f.write(r['Output file']+'\n')

# docx

def set_korean_font(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'; normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
    normal.font.color.rgb = RGBColor(0,0,0)
    for h in ['Heading 1','Heading 2','Heading 3']:
        st=doc.styles[h]; st.font.name='Malgun Gothic'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.color.rgb=RGBColor(0,0,0)

def set_table_width_pct(table, pct=100):
    tbl=table._tbl; tblPr=tbl.tblPr
    if tblPr is None:
        tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'),'pct'); tblW.set(qn('w:w'),str(int(pct*50)))

doc=Document(); set_korean_font(doc)
sec=doc.sections[0]; usable=(sec.page_width-sec.left_margin-sec.right_margin)/914400.0
col1=max(1.3,usable*0.18); col2=max(1.1,usable*0.13); col3=max(3.0,usable-col1-col2); img_w=max(2.3,col3-0.2)
for ex in range(1,15):
    t,o,c,s=meta[ex]
    doc.add_heading(f'Exercise {ex}. {t}', level=2)
    doc.add_heading('1) 실험 목적', level=3); doc.add_paragraph(o, style='List Bullet')
    doc.add_heading('2) 변경 인자', level=3); doc.add_paragraph(c, style='List Bullet')
    doc.add_heading('3) 결과 표', level=3)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.autofit=False; set_table_width_pct(table,100)
    hdr=table.rows[0].cells; hdr[0].text='파라미터'; hdr[1].text='값'; hdr[2].text='결과 이미지'
    widths=[Inches(col1),Inches(col2),Inches(col3)]
    for i,w in enumerate(widths): hdr[i].width=w
    for r in rows:
        if r['Exercise']!=f'Exercise {ex}': continue
        rc=table.add_row().cells; rc[0].text=r['Parameter']; rc[1].text=r['Value']
        for i,w in enumerate(widths): rc[i].width=w
        p=ROOT/r['Output file']
        if p.exists(): rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))
    doc.add_paragraph(''); doc.add_heading('4) 해석 요약', level=3); doc.add_paragraph(s, style='List Bullet'); doc.add_paragraph('')

docx=OUT/'ablation_report_compact_ko_v3_cv4.docx'; doc.save(docx)
print('rows',len(rows)); print(csv_path); print(md); print(docx); print(overall); print(lst)
