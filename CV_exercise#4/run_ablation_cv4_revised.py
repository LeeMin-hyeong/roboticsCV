import csv
from pathlib import Path
import numpy as np
import cv2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.ndimage import grey_dilation, grey_erosion
from skimage.filters import threshold_otsu
from skimage.color import label2rgb
from skimage.measure import label
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#4")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
rows = []


def add_row(ex, param, value, out, analysis, metric=''):
    rows.append({
        'Exercise': f'Exercise {ex}',
        'Parameter': str(param),
        'Value': str(value),
        'Output file': str(out.relative_to(ROOT)).replace('\\', '/'),
        'Visual change': '',
        'Optional metric': metric,
        'One-line analysis': analysis,
    })


def save3(a, b, c, t1, t2, t3, out, cmap='gray'):
    plt.figure(figsize=(12, 4))
    plt.subplot(1, 3, 1); plt.imshow(a, cmap=cmap); plt.title(t1); plt.axis('off')
    plt.subplot(1, 3, 2); plt.imshow(b, cmap=cmap); plt.title(t2); plt.axis('off')
    plt.subplot(1, 3, 3); plt.imshow(c, cmap=cmap); plt.title(t3); plt.axis('off')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()

# Exercise 1: Binary dilation (kernel study)
exd = OUT / 'ex1_binary_dilation'; exd.mkdir(exist_ok=True)
img = cv2.imread(str(ROOT / 'bacteria.png'), cv2.IMREAD_GRAYSCALE)
for k in [10, 20, 30]:
    se = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (k, k))
    d = cv2.dilate(img, se)
    out = exd / f'ex1_disk_{k}.png'
    save3(img, d, cv2.absdiff(d, img), '원본', f'dilation disk={k}', 'difference', out)
    add_row(1, 'disk_size', k, out, f'disk={k}에서 전경이 확장되고 작은 공백이 메워졌다.', f'diff_mean={float(np.mean(cv2.absdiff(d, img))):.2f}')

# Exercise 2: Binary erosion (kernel study)
exd = OUT / 'ex2_binary_erosion'; exd.mkdir(exist_ok=True)
for k in [3, 7, 11]:
    se = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (k, k))
    e = cv2.erode(img, se)
    out = exd / f'ex2_disk_{k}.png'
    save3(img, e, cv2.absdiff(img, e), '원본', f'erosion disk={k}', 'difference', out)
    add_row(2, 'disk_size', k, out, f'disk={k}에서 침식이 강해져 얇은 구조가 감소했다.', f'diff_mean={float(np.mean(cv2.absdiff(img, e))):.2f}')

# Exercise 3: Binary erosion coins (shape/size study in original code)
exd = OUT / 'ex3_binary_erosion_coins'; exd.mkdir(exist_ok=True)
cir = cv2.imread(str(ROOT / 'circles.png'), cv2.IMREAD_GRAYSCALE)
for shape, size in [('square', 30), ('square', 70), ('square', 96), ('disk', 30), ('disk', 70), ('disk', 96)]:
    if shape == 'square':
        se = cv2.getStructuringElement(cv2.MORPH_RECT, (size, size))
    else:
        se = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (size, size))
    e = cv2.erode(cir, se)
    out = exd / f'ex3_{shape}_{size}.png'
    save3(cir, e, cv2.absdiff(cir, e), '원본', f'{shape} {size}', 'difference', out)
    add_row(3, 'SE(shape,size)', f'{shape},{size}', out, f'{shape} {size}에서 동전 분리 양상이 달라졌다.', f'fg_ratio={float(np.mean(e>0)):.4f}')

# Exercise 4: Binary erosion fence (length is core)
exd = OUT / 'ex4_binary_erosion_fence'; exd.mkdir(exist_ok=True)
fence = cv2.imread(str(ROOT / 'fence.jpg'), cv2.IMREAD_GRAYSCALE)
_, bw_f = cv2.threshold(fence, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
for length in [101, 151, 201]:
    se = cv2.getStructuringElement(cv2.MORPH_CROSS, (length, length))
    det = cv2.erode(bw_f, se)
    out = exd / f'ex4_length_{length}.png'
    save3(fence, bw_f, det, '원본', 'binarized', f'eroded length={length}', out)
    add_row(4, 'cross_length', length, out, f'길이 {length}에서 fence 검출 폭이 달라졌다.', f'fg_ratio={float(np.mean(det>0)):.4f}')

# Exercise 5: Small hole removal (keep kernel=10, vary threshold)
exd = OUT / 'ex5_small_hole_removal'; exd.mkdir(exist_ok=True)
peter = cv2.imread(str(ROOT / 'peter.png'), cv2.IMREAD_GRAYSCALE)
se10 = cv2.getStructuringElement(cv2.MORPH_RECT, (10, 10))
for th in [90, 100, 110]:
    _, mask = cv2.threshold(peter, th, 255, cv2.THRESH_BINARY_INV)
    dil = cv2.dilate(mask, se10)
    clo = cv2.erode(dil, se10)
    diff = cv2.subtract(clo, mask)
    out = exd / f'ex5_threshold_{th}.png'
    save3(mask, clo, diff, f'binary t={th}', 'closed', 'filled holes', out)
    add_row(5, 'threshold', th, out, f'threshold={th}에서 hole 채움 대상 영역이 변했다.', f'filled_ratio={float(np.mean(diff>0)):.4f}')

# Exercise 6: Binary edge detection (kernel fixed, mode study)
exd = OUT / 'ex6_binary_edge_detection'; exd.mkdir(exist_ok=True)
clip = cv2.imread(str(ROOT / 'cliparts.png'), cv2.IMREAD_GRAYSCALE)
se9 = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (9, 9))
dil = cv2.dilate(clip, se9)
ero = cv2.erode(clip, se9)
edge1 = cv2.subtract(dil, clip)
edge2 = cv2.subtract(clip, ero)
edge3 = cv2.add(edge1, edge2)
for mode, edge in [('edge1_dilated_minus_orig', edge1), ('edge2_orig_minus_eroded', edge2), ('edge3_combined', edge3)]:
    out = exd / f'ex6_{mode}.png'
    save3(clip, edge, cv2.threshold(edge, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1], '원본', mode, 'binarized edge', out)
    add_row(6, 'edge_mode', mode, out, f'{mode} 방식에서 경계 응답 분포가 달랐다.', f'edge_mean={float(edge.mean()):.2f}')

# Exercise 7: 1D gray morph (original loop)
exd = OUT / 'ex7_1d_gray_morph'; exd.mkdir(exist_ok=True)
imgp = cv2.imread(str(ROOT / 'peter.png'), cv2.IMREAD_GRAYSCALE)
sig = imgp[150:249, 180]
for length in [1, 11, 21]:
    se = np.ones((length,), dtype=np.uint8)
    d = grey_dilation(sig, footprint=se)
    e = grey_erosion(sig, footprint=se)
    out = exd / f'ex7_length_{length}.png'
    plt.figure(figsize=(10,4))
    plt.plot(sig, 'bo-', label='orig'); plt.plot(d, 'r-', label='dil'); plt.plot(e, 'g-', label='ero')
    plt.title(f'length={length}'); plt.legend(); plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
    add_row(7, 'SE_length', length, out, f'길이 {length}에서 1D 신호 envelope 폭이 달라졌다.', f'd_range={float(d.max()-d.min()):.2f}')

# Exercise 8: 2D gray morph (structuring element type in original code)
exd = OUT / 'ex8_2d_gray_morph'; exd.mkdir(exist_ok=True)
but = cv2.imread(str(ROOT / 'butterfly.png'))
se_defs = {
    'square_10': cv2.getStructuringElement(cv2.MORPH_RECT, (10,10)),
    'disk_17': cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (17,17)),
    'double_line': np.pad(np.ones((2,18), dtype=np.uint8), ((0,16),(0,0))),
}
for name, se in se_defs.items():
    d = cv2.dilate(but, se)
    out = exd / f'ex8_{name}.png'
    save3(cv2.cvtColor(but, cv2.COLOR_BGR2RGB), cv2.cvtColor(d, cv2.COLOR_BGR2RGB), cv2.cvtColor(cv2.absdiff(d, but), cv2.COLOR_BGR2RGB), 'orig', name, 'difference', out, cmap=None)
    add_row(8, 'SE_type', name, out, f'{name} 구조요소에서 강조 방향과 연결 특성이 달라졌다.', f'diff_mean={float(np.mean(cv2.absdiff(d,but))):.2f}')

# Exercise 9: Coin separation by gray dilation (kernel is core)
exd = OUT / 'ex9_coin_separation'; exd.mkdir(exist_ok=True)
coins = cv2.imread(str(ROOT / 'coins.png'), cv2.IMREAD_GRAYSCALE)
for size in [41, 61, 81]:
    se = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (size, size))
    dil = cv2.dilate(coins, se)
    bw = cv2.bitwise_not((dil > threshold_otsu(dil)).astype(np.uint8))
    L = label(bw, connectivity=2)
    rgb = label2rgb(L, bg_label=0)
    out = exd / f'ex9_disk_{size}.png'
    save3(coins, dil, rgb, 'orig', f'dilated {size}', 'labeled', out, cmap='gray')
    add_row(9, 'disk_size', size, out, f'disk {size}에서 동전 결합/분리 상태가 달라졌다.', f'num_labels={int(L.max())}')

# Exercise 10: Hole detection by gray erosion (length is core)
exd = OUT / 'ex10_hole_detection'; exd.mkdir(exist_ok=True)
fencef = fence.astype(np.float32) / 255.0
for length in [101, 151, 201]:
    nhood = np.zeros((length, length), dtype=np.uint8)
    nhood[length//2, :] = 1; nhood[:, length//2] = 1
    er = cv2.erode(fencef, nhood)
    bw = er > threshold_otsu(er)
    out = exd / f'ex10_length_{length}.png'
    save3(fencef, er, bw, 'orig', f'eroded {length}', 'thresholded', out)
    add_row(10, 'cross_length', length, out, f'길이 {length}에서 hole 후보 응답 강도가 달라졌다.', f'fg_ratio={float(np.mean(bw)):.4f}')

# Exercise 11: Gray morphological edge detector (keep kernel=7, vary threshold)
exd = OUT / 'ex11_gray_edge_detector'; exd.mkdir(exist_ok=True)
bike = cv2.imread(str(ROOT / 'bike.png'), cv2.IMREAD_GRAYSCALE).astype(np.float32)/255.0
se7 = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (7,7))
dil = cv2.dilate(bike, se7)
edge = dil - bike
for th in [0.10, 0.15, 0.20]:
    bw = edge > th
    out = exd / f'ex11_threshold_{th:.2f}.png'
    save3(bike, edge, bw, 'orig', 'edge map', f'edge>{th:.2f}', out)
    add_row(11, 'edge_threshold', f'{th:.2f}', out, f'threshold {th:.2f}에서 경계 선택 강도가 달라졌다.', f'fg_ratio={float(np.mean(bw)):.4f}')

# Exercise 12: Cascaded gray dilations (iterations are core)
exd = OUT / 'ex12_cascaded_dilations'; exd.mkdir(exist_ok=True)
but2 = cv2.imread(str(ROOT / 'butterfly.png'))
se1 = cv2.getStructuringElement(cv2.MORPH_RECT, (10,1))
se2 = cv2.warpAffine(cv2.getStructuringElement(cv2.MORPH_RECT, (10,1)), cv2.getRotationMatrix2D((5,5),45,1.0), (10,10))
se3 = cv2.warpAffine(cv2.getStructuringElement(cv2.MORPH_RECT, (10,1)), cv2.getRotationMatrix2D((5,5),135,1.0), (10,10))
d1 = cv2.dilate(but2, se1); d2 = cv2.dilate(d1, se2); d3 = cv2.dilate(d2, se3)
for it, im in [(1,d1),(2,d2),(3,d3)]:
    out = exd / f'ex12_iter_{it}.png'
    save3(cv2.cvtColor(but2, cv2.COLOR_BGR2RGB), cv2.cvtColor(im, cv2.COLOR_BGR2RGB), cv2.cvtColor(cv2.absdiff(im,but2), cv2.COLOR_BGR2RGB), 'orig', f'iter={it}', 'difference', out, cmap=None)
    add_row(12, 'iterations', it, out, f'cascaded dilation {it}단계에서 방향성 팽창이 누적됐다.', f'diff_mean={float(np.mean(cv2.absdiff(im,but2))):.2f}')

# Exercise 13: Majority/Median filter (kernel variation already in original)
exd = OUT / 'ex13_majority_median'; exd.mkdir(exist_ok=True)
bac = cv2.imread(str(ROOT / 'bacteria.png'), cv2.IMREAD_GRAYSCALE)
np.random.seed(42)
mask1 = np.random.choice([0,1,2], size=bac.shape, p=[0.95,0.025,0.025])
no1 = bac.copy(); no1[mask1==1]=0; no1[mask1==2]=255
mask2 = np.random.choice([0,1,2], size=bac.shape, p=[0.8,0.1,0.1])
no2 = no1.copy(); no2[mask2==1]=0; no2[mask2==2]=255
for level_name, noisy in [('noise5', no1), ('noise20', no2)]:
    den = cv2.medianBlur(noisy, 3)
    out = exd / f'ex13_{level_name}_k3.png'
    save3(bac, noisy, den, 'orig', level_name, 'median k3', out)
    add_row(13, 'noise_level', level_name, out, f'{level_name}에서 median 3x3이 salt-pepper 잡음을 완화했다.', f'mae={float(np.mean(np.abs(den.astype(np.float32)-bac.astype(np.float32)))):.2f}')

sc = cv2.imread(str(ROOT / 'sculpture.png'))
maskc = np.random.choice([0,1,2], size=sc.shape[:2], p=[0.95,0.025,0.025])
no = sc.copy()
for ch in range(3):
    layer = no[:,:,ch]
    layer[maskc==1]=0; layer[maskc==2]=255
    no[:,:,ch]=layer
for k in [3,7]:
    den = cv2.medianBlur(no, k)
    out = exd / f'ex13_sculpture_k{k}.png'
    save3(cv2.cvtColor(sc,cv2.COLOR_BGR2RGB), cv2.cvtColor(no,cv2.COLOR_BGR2RGB), cv2.cvtColor(den,cv2.COLOR_BGR2RGB), 'orig', 'noisy', f'median {k}x{k}', out, cmap=None)
    add_row(13, 'median_kernel', k, out, f'sculpture에서 median {k}x{k}로 잡음이 감소했다.', f'diff_mean={float(np.mean(cv2.absdiff(den,sc))):.2f}')

# Exercise 14: Nonuniform lighting compensation (window size is core)
exd = OUT / 'ex14_lighting_comp'; exd.mkdir(exist_ok=True)
paper = cv2.imread(str(ROOT / 'paper.png'), cv2.IMREAD_GRAYSCALE)
imgf = paper.astype(np.float32)/255.0
for w in [31, 61, 91]:
    rankf = cv2.blur(paper, (w,w)).astype(np.float32)/255.0
    dif = rankf - imgf
    bw = dif > threshold_otsu(dif)
    out = exd / f'ex14_window_{w}.png'
    save3(imgf, dif, bw, 'orig', f'filtered-orig (w={w})', 'thresholded', out)
    add_row(14, 'window_size', w, out, f'window {w}에서 조명 보정 강도와 문자 분리가 달라졌다.', f'fg_ratio={float(np.mean(bw)):.4f}')

# save files
csv_path = OUT / 'ablation_summary_cv4_v2.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w = csv.DictWriter(f, fieldnames=['Exercise','Parameter','Value','Output file','Visual change','Optional metric','One-line analysis'])
    w.writeheader(); w.writerows(rows)

overall = OUT / 'overall_summary_table_cv4_v2.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 4, revised)\n\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows:
        f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta = {
1:('Binary_dilation','이진 dilation 크기에 따른 연결 효과를 비교한다.','disk_size','원 코드처럼 disk 크기를 키우면 전경이 확장됐다. 작은 결손 복원에는 유리했다. 큰 disk에서는 객체 병합이 쉽게 발생했다.'),
2:('Binary_Erosion','이진 erosion 크기 변화에 따른 객체 축소를 확인한다.','disk_size','원 코드의 침식 실험에서 disk가 커질수록 구조가 빠르게 감소했다. 얇은 전경은 먼저 사라졌다. 잡음 억제와 정보 보존 사이 절충이 필요했다.'),
3:('Binary_Erosion_Coins','동전 분리를 위한 structuring element 형태와 크기를 비교한다.','SE(shape,size)','square와 disk는 같은 크기에서도 분리 양상이 달랐다. 크기를 키우면 접촉 영역은 줄었다. 과도한 침식은 객체 자체를 약화시켰다.'),
4:('Binary_Erosion_Fence','cross 길이에 따른 fence 검출 변화를 비교한다.','cross_length','길이가 길수록 구조요소 방향성이 강해졌다. fence 패턴 강조는 증가했다. 동시에 검출 영역은 더 보수적으로 변했다.'),
5:('Small_Hole_Removal','small hole removal에서 이진화 임계값 영향을 비교한다.','threshold','커널은 원 코드 값으로 고정했다. threshold 변화만으로 hole 후보 영역이 달라졌다. 채움 결과도 임계값에 민감하게 변했다.'),
6:('Binary_Edge_Detection','binary edge detection에서 edge 조합 방식 차이를 비교한다.','edge_mode','커널은 원 코드 값으로 고정했다. edge1과 edge2는 강조 방향이 다르다. combined 결과는 경계 범위가 가장 넓게 나타났다.'),
7:('1D Gray Morph','1D 회색조 morphology에서 창 길이 효과를 비교한다.','SE_length','원 코드 길이 증가에 따라 envelope 폭이 커졌다. peak와 valley 왜곡도 함께 증가했다. 길이 선택이 신호 형태 보존에 중요했다.'),
8:('2D Gray Morph','2D 회색조 morphology에서 구조요소 형태 차이를 비교한다.','SE_type','shape가 달라지면 강조 방향이 달라졌다. line 계열은 방향성 반응이 강했다. disk와 square는 보다 균일한 팽창 특성을 보였다.'),
9:('Coin Separation by Graylevel Dilation','gray dilation 크기에 따른 coin 분리 변화를 비교한다.','disk_size','dilation 크기에 따라 동전 blob 결합 정도가 바뀌었다. 분리 성능은 크기 선택에 민감했다. 과도한 크기는 객체 결합을 유발했다.'),
10:('Hole Detection by Graylevel Erosion','cross 길이에 따른 hole 강조 효과를 비교한다.','cross_length','길이가 커질수록 침식 기반 hole 응답이 강화됐다. threshold 결과도 함께 변했다. 지나치게 긴 길이는 배경 영향까지 확대했다.'),
11:('Graylevel Morphological Edge Detector','edge threshold에 따른 경계 선택 강도를 비교한다.','edge_threshold','커널은 원 코드 값으로 고정했다. threshold를 높이면 강한 경계만 남았다. 낮은 threshold에서는 배경 잡응답이 늘었다.'),
12:('Cascaded Graylevel Dilations','cascaded dilation 단계별 누적 효과를 비교한다.','iterations','단계가 늘수록 방향성 팽창이 누적됐다. 작은 틈은 메워졌다. 세부 경계는 점차 완화됐다.'),
13:('Majority Filter and Median Filter','원 코드 노이즈 조건에서 median filtering 효과를 비교한다.','noise_level / median_kernel','노이즈 강도가 높을수록 원본 손상이 커졌다. median 필터는 잡음 감소에 효과적이었다. 큰 커널은 노이즈를 더 줄이지만 텍스처도 완화됐다.'),
14:('Nonuniform Lighting Compensation','조명 보정 window 크기 변화 효과를 비교한다.','window_size','window가 커질수록 저주파 조명 성분 제거가 강해졌다. 문자 대비는 개선될 수 있다. 과도한 window는 세부 정보 약화를 유발했다.'),
}

md = OUT / 'ablation_report_compact_ko_v3_cv4_revised.md'
with md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 4 Ablation Study (Revised)\n\n')
    for ex in range(1,15):
        t,o,c,s=meta[ex]
        f.write(f'## Exercise {ex}. {t}\n\n### 1) 실험 목적\n- {o}\n\n### 2) 변경 인자\n- {c}\n\n### 3) 결과 표\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise']==f'Exercise {ex}':
                f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write(f"\n### 4) 해석 요약\n- {s}\n\n")

lst = OUT / 'generated_images_cv4_v2.txt'
with lst.open('w', encoding='utf-8') as f:
    for r in rows:
        f.write(r['Output file'] + '\n')

# docx

def set_korean_font(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal.font.color.rgb = RGBColor(0,0,0)
    for h in ['Heading 1','Heading 2','Heading 3']:
        st=doc.styles[h]
        st.font.name='Malgun Gothic'
        st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
        st.font.color.rgb = RGBColor(0,0,0)


def set_table_width_pct(table, pct=100):
    tbl=table._tbl
    tblPr=tbl.tblPr
    if tblPr is None:
        tblPr=OxmlElement('w:tblPr')
        tbl.insert(0,tblPr)
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW=OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'),'pct')
    tblW.set(qn('w:w'),str(int(pct*50)))


doc=Document(); set_korean_font(doc)
sec=doc.sections[0]
usable=(sec.page_width-sec.left_margin-sec.right_margin)/914400.0
col1=max(1.3,usable*0.18); col2=max(1.1,usable*0.13); col3=max(3.0,usable-col1-col2); img_w=max(2.3,col3-0.2)

for ex in range(1,15):
    t,o,c,s=meta[ex]
    doc.add_heading(f'Exercise {ex}. {t}', level=2)
    doc.add_heading('1) 실험 목적', level=3); doc.add_paragraph(o, style='List Bullet')
    doc.add_heading('2) 변경 인자', level=3); doc.add_paragraph(c, style='List Bullet')
    doc.add_heading('3) 결과 표', level=3)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.autofit=False; set_table_width_pct(table,100)
    hdr=table.rows[0].cells; hdr[0].text='파라미터'; hdr[1].text='값'; hdr[2].text='결과 이미지'
    widths=[Inches(col1), Inches(col2), Inches(col3)]
    for i,w in enumerate(widths): hdr[i].width=w
    for r in rows:
        if r['Exercise'] != f'Exercise {ex}':
            continue
        rc=table.add_row().cells
        rc[0].text=r['Parameter']; rc[1].text=r['Value']
        for i,w in enumerate(widths): rc[i].width=w
        p=ROOT/r['Output file']
        if p.exists(): rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))
    doc.add_paragraph('')
    doc.add_heading('4) 해석 요약', level=3)
    doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('')


docx = OUT / 'ablation_report_compact_ko_v3_cv4_revised.docx'
doc.save(docx)

print('rows', len(rows))
print(csv_path)
print(md)
print(docx)
print(overall)
print(lst)
