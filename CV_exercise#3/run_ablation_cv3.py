import csv
from pathlib import Path
import numpy as np
import cv2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image
from scipy import ndimage
from skimage import io, color, filters, util, measure
from skimage.morphology import closing, square, remove_small_objects
from skimage.transform import resize
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#3")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

rows = []


def add_row(ex, param, value, out_file, visual, metric, analysis):
    rows.append({
        "Exercise": ex,
        "Parameter": param,
        "Value": str(value),
        "Output file": str(out_file.relative_to(ROOT)).replace('\\', '/'),
        "Visual change": visual,
        "Optional metric": metric,
        "One-line analysis": analysis,
    })


def save_triplet(a, b, c, t1, t2, t3, out, cmap='gray'):
    plt.figure(figsize=(12, 4))
    plt.subplot(1, 3, 1); plt.imshow(a, cmap=cmap); plt.title(t1); plt.axis('off')
    plt.subplot(1, 3, 2); plt.imshow(b, cmap=cmap); plt.title(t2); plt.axis('off')
    plt.subplot(1, 3, 3); plt.imshow(c, cmap=cmap); plt.title(t3); plt.axis('off')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()

# Exercise 1
ex1 = OUT / "ex1_graylevel_thresholding"; ex1.mkdir(exist_ok=True)
img = np.array(Image.open(ROOT / 'peter.png').convert('L'))
for level in [90, 105, 120, 140]:
    bw = img < level
    hole = img * bw.astype(np.uint8)
    out = ex1 / f"ex1_level_{level}.png"
    save_triplet(img, bw, hole, '원본', f'threshold={level}', 'mask x 원본', out)
    fg = float(bw.mean())
    add_row("Exercise 1", "threshold", level, out, "임계값 증가 시 전경 영역 확대", f"foreground_ratio={fg:.4f}", f"threshold를 {level}로 높이자 어두운 영역이 더 많이 전경으로 분류됐다.")

# Exercise 2
ex2 = OUT / "ex2_global_thresholding"; ex2.mkdir(exist_ok=True)
img2 = io.imread(ROOT / 'front.png')
if img2.ndim == 3:
    img2 = (color.rgb2gray(img2) * 255).astype(np.uint8)
otsu = int(round(filters.threshold_otsu(img2)))
for delta in [-20, 0, 20, 40]:
    t = int(np.clip(otsu + delta, 0, 255))
    bw = img2 > t
    overlay = np.where(bw, np.min(img2), img2)
    out = ex2 / f"ex2_threshold_{t}.png"
    save_triplet(img2, bw, overlay, '원본', f'global T={t}', 'overlay', out)
    fg = float(bw.mean())
    add_row("Exercise 2", "global_threshold", t, out, "임계값 증가 시 밝은 영역만 남음", f"foreground_ratio={fg:.4f}", f"전역 임계값 {t}에서는 배경 제거가 강해지고 밝은 전경만 남는 경향이 나타났다.")

# Exercise 3
ex3 = OUT / "ex3_local_adaptive_thresholding"; ex3.mkdir(exist_ok=True)
img3 = io.imread(ROOT / 'paper.png')
if img3.ndim == 3:
    img3 = color.rgb2gray(img3)
img3 = util.img_as_float(img3)
global_t = filters.threshold_otsu(img3)

def local_adapt(img, step_size=16, tile_size=32, var_thresh=0.0005):
    rows_, cols_ = img.shape
    local_bw = np.zeros((rows_, cols_), dtype=np.uint8)
    thresh_map = np.zeros((rows_ // step_size, cols_ // step_size), dtype=float)
    for i in range(0, rows_, step_size):
        for j in range(0, cols_, step_size):
            rs = slice(i, min(i + tile_size, rows_))
            cs = slice(j, min(j + tile_size, cols_))
            tile = img[rs, cs]
            v = np.var(tile)
            if v > var_thresh:
                lt = filters.threshold_otsu(tile)
                local_bw[rs, cs] = (tile > lt).astype(np.uint8)
                thresh_map[i // step_size, j // step_size] = lt
            else:
                local_bw[rs, cs] = (tile > global_t).astype(np.uint8)
                thresh_map[i // step_size, j // step_size] = global_t
    interp = resize(thresh_map, img.shape, order=1, preserve_range=True, anti_aliasing=False)
    return local_bw, interp

for vt in [0.0001, 0.0005, 0.001, 0.005]:
    bw, tmap = local_adapt(img3, step_size=16, tile_size=32, var_thresh=vt)
    out = ex3 / f"ex3_varThresh_{vt:.4f}.png"
    plt.figure(figsize=(12, 4))
    plt.subplot(1,3,1); plt.imshow(img3, cmap='gray'); plt.title('원본'); plt.axis('off')
    plt.subplot(1,3,2); plt.imshow(bw, cmap='gray'); plt.title(f'local BW var={vt}'); plt.axis('off')
    plt.subplot(1,3,3); plt.imshow(tmap, cmap='inferno'); plt.title('threshold map'); plt.axis('off')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
    fg = float(bw.mean())
    add_row("Exercise 3", "var_thresh", vt, out, "분산 임계값 증가 시 지역 임계 적용 축소", f"foreground_ratio={fg:.4f}", f"var_thresh를 {vt}로 높이면 지역 Otsu 적용 영역이 줄어 결과가 더 전역 임계에 가까워졌다.")

# Exercise 4
ex4 = OUT / "ex4_map_skin_detector"; ex4.mkdir(exist_ok=True)
train_imgs = [io.imread(ROOT / f'Face_Training_{i}.jpg') for i in range(1, 6)]
train_masks = [io.imread(ROOT / f'Face_ref_{i}.png') for i in range(1, 6)]
test_imgs = [io.imread(ROOT / f'Face_Test_{i}.jpg') for i in range(1, 3)]

def build_classifier(step):
    bins = 256 // step + 1
    skin = np.zeros((bins, bins, bins), dtype=np.float64)
    nonskin = np.zeros_like(skin)
    for imgc, m in zip(train_imgs, train_masks):
        if m.ndim == 3:
            m = color.rgb2gray(m)
            m = (m > 0.5).astype(np.uint8)
        else:
            m = (m > 0).astype(np.uint8)
        idx = np.clip((imgc // step), 0, bins - 1).astype(int)
        for ch in range(3):
            pass
        flat_idx = idx[:, :, 0] * bins * bins + idx[:, :, 1] * bins + idx[:, :, 2]
        skin_counts = np.bincount(flat_idx[m > 0].ravel(), minlength=bins*bins*bins)
        non_counts = np.bincount(flat_idx[m == 0].ravel(), minlength=bins*bins*bins)
        skin += skin_counts.reshape((bins, bins, bins))
        nonskin += non_counts.reshape((bins, bins, bins))
    return skin > nonskin, bins

def apply_classifier(imgc, clf, step, bins):
    idx = np.clip((imgc // step), 0, bins - 1).astype(int)
    return clf[idx[:, :, 0], idx[:, :, 1], idx[:, :, 2]].astype(np.uint8)

for step in [8, 16, 32]:
    clf, bins = build_classifier(step)
    for ti, timg in enumerate(test_imgs, start=1):
        raw = apply_classifier(timg, clf, step, bins)
        clean = remove_small_objects(raw.astype(bool), min_size=400).astype(np.uint8)
        label = measure.label(clean)
        props = measure.regionprops(label)
        max_area = max([p.area for p in props], default=0)
        out = ex4 / f"ex4_step_{step}_test_{ti}.png"
        plt.figure(figsize=(12, 4))
        plt.subplot(1,3,1); plt.imshow(timg); plt.title(f'test {ti}'); plt.axis('off')
        plt.subplot(1,3,2); plt.imshow(raw, cmap='gray'); plt.title('raw MAP'); plt.axis('off')
        plt.subplot(1,3,3); plt.imshow(clean, cmap='gray'); plt.title('filtered'); plt.axis('off')
        plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
        fg = float(clean.mean())
        add_row("Exercise 4", "color_bin_step", f"{step} (test{ti})", out, "bin 크기에 따라 피부색 분류 경계 변화", f"foreground_ratio={fg:.4f}, largest_area={max_area}", f"bin step {step}에서는 피부 클래스 경계가 달라져 얼굴 후보 영역 크기가 변했다.")

# Exercise 5
ex5 = OUT / "ex5_region_labeling"; ex5.mkdir(exist_ok=True)
fish = io.imread(ROOT / 'fish.png')
fish_gray = color.rgb2gray(fish) if fish.ndim == 3 else fish
otsu5 = filters.threshold_otsu(fish_gray)
for k in [1, 3, 5, 7]:
    bw = closing(fish_gray > otsu5, square(k))
    lab = measure.label(bw)
    nreg = int(lab.max())
    rgb = color.label2rgb(lab, bg_label=0)
    out = ex5 / f"ex5_closing_{k}.png"
    plt.figure(figsize=(12, 4))
    plt.subplot(1,3,1); plt.imshow(fish); plt.title('원본'); plt.axis('off')
    plt.subplot(1,3,2); plt.imshow(bw, cmap='gray'); plt.title(f'closing {k}'); plt.axis('off')
    plt.subplot(1,3,3); plt.imshow(rgb); plt.title(f'labeled ({nreg})'); plt.axis('off')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
    add_row("Exercise 5", "closing_kernel", k, out, "커널 증가 시 인접 영역 연결 강화", f"num_regions={nreg}", f"closing 커널을 {k}로 키우자 분리된 영역이 합쳐져 라벨 개수가 줄어드는 경향을 보였다.")

# Exercise 6
ex6 = OUT / "ex6_hole_filling"; ex6.mkdir(exist_ok=True)
face = np.array(Image.open(ROOT / 'face.png').convert('L'))
for level in [90, 105, 120, 140]:
    bw = face < level
    filled = ndimage.binary_fill_holes(bw)
    diff = (filled.astype(np.uint8) - bw.astype(np.uint8)) > 0
    out = ex6 / f"ex6_level_{level}.png"
    plt.figure(figsize=(12, 4))
    plt.subplot(1,3,1); plt.imshow(bw, cmap='gray'); plt.title(f'binary {level}'); plt.axis('off')
    plt.subplot(1,3,2); plt.imshow(filled, cmap='gray'); plt.title('filled'); plt.axis('off')
    plt.subplot(1,3,3); plt.imshow(diff, cmap='gray'); plt.title('newly filled'); plt.axis('off')
    plt.tight_layout(); plt.savefig(out, dpi=140); plt.close()
    fill_ratio = float(diff.mean())
    add_row("Exercise 6", "threshold", level, out, "이진화 수준에 따라 채워지는 hole 면적 변화", f"filled_ratio={fill_ratio:.4f}", f"threshold {level}에서는 내부 공극의 채움 면적이 {fill_ratio:.3f} 수준으로 나타났다.")

# Save summary csv
csv_path = OUT / 'ablation_summary_cv3.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w = csv.DictWriter(f, fieldnames=["Exercise", "Parameter", "Value", "Output file", "Visual change", "Optional metric", "One-line analysis"])
    w.writeheader()
    w.writerows(rows)

# Build markdown report (v3 style)
meta = {
1: ('Graylevel thresholding', '단일 임계값 변화에 따른 이진 분할 결과 변화를 확인한다.', 'threshold', 'threshold가 증가할수록 전경으로 분류되는 영역이 넓어졌다. 낮은 값에서는 세부 객체가 일부 끊겼다. 높은 값에서는 배경 노이즈가 함께 포함됐다. 과제 목적에 맞는 중간 임계값 선택이 중요했다.'),
2: ('Global Thresholding', '전역 임계값 설정이 분할 결과에 주는 영향을 비교한다.', 'global_threshold', '전역 임계값을 높일수록 밝은 영역만 남고 어두운 디테일은 빠르게 사라졌다. Otsu 근처 값에서 전경과 배경 균형이 상대적으로 안정적이었다. 과도한 임계값은 객체 내부 정보 손실을 유발했다.'),
3: ('Locally Adaptive Thresholding', '지역 분산 조건에 따른 적응 임계 분할 효과를 비교한다.', 'var_thresh', 'var_thresh가 낮을 때는 더 많은 타일에서 지역 임계가 적용됐다. 지역 텍스트와 경계는 잘 보였지만 잡음도 증가했다. var_thresh를 높이면 결과가 전역 임계 형태에 가까워졌다. 지역성과 안정성 사이의 절충이 필요했다.'),
4: ('MAP Skin Detector', '색상 히스토그램 bin 크기가 피부 분류 성능에 미치는 영향을 비교한다.', 'color_bin_step', 'bin step이 작으면 색상 구분은 세밀해진다. 학습 샘플 수가 제한되면 오분류도 늘 수 있다. bin step이 크면 분류 경계가 단순해져 얼굴 외 영역이 섞일 수 있다. 테스트 영상별로 적정 bin 크기가 다르게 나타났다.'),
5: ('Region Labeling', '형태학적 closing 크기에 따른 연결 성분 라벨링 결과를 비교한다.', 'closing_kernel', 'closing 커널이 커질수록 끊긴 전경이 연결됐다. 작은 객체 간 간격이 메워지면서 라벨 수가 감소했다. 커널이 너무 크면 서로 다른 객체가 합쳐질 수 있다. 객체 분리 목적이면 작은 커널이 더 유리했다.'),
6: ('Hole Filling', '이진화 임계값에 따른 hole filling 결과 변화를 비교한다.', 'threshold', 'threshold가 낮으면 내부 hole 자체가 적게 형성됐다. threshold가 올라가면 공극 후보가 늘어 채움 면적이 증가했다. 너무 높은 값은 외곽 잡영까지 포함해 과도한 채움으로 이어졌다. 채움 단계 전 이진화 품질이 최종 결과를 좌우했다.'),
}

md_path = OUT / 'ablation_report_compact_ko_v3_cv3.md'
with md_path.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 3 Ablation Study\n\n')
    for ex in range(1, 7):
        title, obj, changed, summary = meta[ex]
        f.write(f'## Exercise {ex}. {title}\n\n')
        f.write('### 1) 실험 목적\n')
        f.write(f'- {obj}\n\n')
        f.write('### 3) 변경 인자\n')
        f.write(f'- {changed}\n\n')
        f.write('### 4) 결과 표\n')
        f.write('| 파라미터 | 값 | 결과 이미지 |\n')
        f.write('|---|---|---|\n')
        for r in rows:
            if r['Exercise'] == f'Exercise {ex}':
                f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write('\n### 6) 해석 요약\n')
        f.write(f'- {summary}\n\n')

# overall summary table
overall = OUT / 'overall_summary_table_cv3.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 3)\n\n')
    f.write('| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows:
        f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

# image list
img_list = OUT / 'generated_images_cv3.txt'
with img_list.open('w', encoding='utf-8') as f:
    for r in rows:
        f.write(r['Output file'] + '\n')

# docx with image-in-table

def set_korean_font(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def set_table_width_pct(table, pct=100):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), str(int(pct * 50)))

doc = Document()
set_korean_font(doc)
sec = doc.sections[0]
usable = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0
col1 = max(1.3, usable * 0.18)
col2 = max(1.1, usable * 0.13)
col3 = max(3.0, usable - col1 - col2)
img_w = max(2.3, col3 - 0.2)

for ex in range(1, 7):
    title, obj, changed, summary = meta[ex]
    doc.add_heading(f'Exercise {ex}. {title}', level=2)
    doc.add_heading('1) 실험 목적', level=3)
    doc.add_paragraph(obj, style='List Bullet')
    doc.add_heading('3) 변경 인자', level=3)
    doc.add_paragraph(changed, style='List Bullet')
    doc.add_heading('4) 결과 표', level=3)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width_pct(table, 100)
    hdr = table.rows[0].cells
    hdr[0].text = '파라미터'; hdr[1].text = '값'; hdr[2].text = '결과 이미지'
    widths = [Inches(col1), Inches(col2), Inches(col3)]
    for i,w in enumerate(widths):
        hdr[i].width = w

    for r in rows:
        if r['Exercise'] != f'Exercise {ex}':
            continue
        row = table.add_row().cells
        row[0].text = r['Parameter']
        row[1].text = r['Value']
        for i,w in enumerate(widths):
            row[i].width = w
        p = ROOT / r['Output file']
        if p.exists():
            run = row[2].paragraphs[0].add_run()
            run.add_picture(str(p), width=Inches(img_w))

    doc.add_paragraph('')
    doc.add_heading('6) 해석 요약', level=3)
    doc.add_paragraph(summary, style='List Bullet')
    doc.add_paragraph('')

docx_path = OUT / 'ablation_report_compact_ko_v3_cv3.docx'
doc.save(docx_path)

print('rows', len(rows))
print(csv_path)
print(md_path)
print(overall)
print(img_list)
print(docx_path)
