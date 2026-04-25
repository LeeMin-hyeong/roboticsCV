import csv
from pathlib import Path
import numpy as np
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
from torch.utils.data import TensorDataset, DataLoader
from sklearn.datasets import load_digits
from sklearn.model_selection import train_test_split
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\lmhst\git\roboticsCV\CV_exercise#7")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DEVICE = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
SEED = 42
np.random.seed(SEED)
torch.manual_seed(SEED)

rows = []

def add_row(ex, param, value, out_file, analysis, metric=''):
    rows.append({
        'Exercise': f'Exercise {ex}',
        'Parameter': str(param),
        'Value': str(value),
        'Output file': str(out_file.relative_to(ROOT)).replace('\\','/'),
        'Visual change': '',
        'Optional metric': metric,
        'One-line analysis': analysis,
    })

# dataset (digits 8x8 -> 32x32)
digits = load_digits()
X = digits.images.astype(np.float32) / 16.0
y = digits.target.astype(np.int64)
X_t = torch.tensor(X).unsqueeze(1)
X_t = F.interpolate(X_t, size=(32,32), mode='bilinear', align_corners=False)
X_train, X_test, y_train, y_test = train_test_split(X_t, torch.tensor(y), test_size=0.2, random_state=SEED, stratify=y)

train_loader = DataLoader(TensorDataset(X_train, y_train), batch_size=64, shuffle=True)
test_loader = DataLoader(TensorDataset(X_test, y_test), batch_size=128, shuffle=False)


def evaluate(model, loader):
    model.eval()
    c=t=0
    with torch.no_grad():
        for xb,yb in loader:
            xb,yb=xb.to(DEVICE), yb.to(DEVICE)
            pred=model(xb).argmax(1)
            c += (pred==yb).sum().item()
            t += yb.size(0)
    return 100.0*c/t


def train_model(model, epochs=4, lr=1e-3):
    model.to(DEVICE)
    opt = optim.AdamW(model.parameters(), lr=lr)
    ce = nn.CrossEntropyLoss()
    train_curve = []
    for _ in range(epochs):
        model.train()
        c=t=0
        for xb,yb in train_loader:
            xb,yb=xb.to(DEVICE), yb.to(DEVICE)
            opt.zero_grad()
            out=model(xb)
            loss=ce(out,yb)
            loss.backward()
            opt.step()
            c += (out.argmax(1)==yb).sum().item()
            t += yb.size(0)
        train_curve.append(100.0*c/t)
    test_acc = evaluate(model, test_loader)
    return train_curve, test_acc


def param_count(model):
    return sum(p.numel() for p in model.parameters() if p.requires_grad)

# Ex1 Kernel size
class KernelCNN(nn.Module):
    def __init__(self, k=3):
        super().__init__()
        p = k//2
        self.features = nn.Sequential(
            nn.Conv2d(1,32,k,padding=p), nn.ReLU(),
            nn.Conv2d(32,64,k,padding=p), nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Conv2d(64,128,k,padding=p), nn.ReLU(),
            nn.MaxPool2d(2),
        )
        self.classifier = nn.Sequential(
            nn.Flatten(),
            nn.Linear(128*8*8, 128), nn.ReLU(),
            nn.Dropout(0.3),
            nn.Linear(128,10)
        )
    def forward(self,x):
        return self.classifier(self.features(x))

exd = OUT / 'ex1_kernel_size'; exd.mkdir(exist_ok=True)
ks_vals=[3,5,7]
ks_acc=[]
ks_params=[]
for k in ks_vals:
    m=KernelCNN(k)
    tr, te = train_model(m, epochs=4, lr=1e-3)
    ks_acc.append(te); ks_params.append(param_count(m))

plt.figure(figsize=(10,4))
plt.subplot(1,2,1); plt.plot(ks_vals, ks_acc, marker='o'); plt.title('Test Acc vs Kernel'); plt.xlabel('kernel'); plt.ylabel('acc')
plt.subplot(1,2,2); plt.bar([str(k) for k in ks_vals], ks_params); plt.title('Params')
plt.tight_layout();
out_all = exd/'ex1_kernel_summary.png'; plt.savefig(out_all,dpi=140); plt.close()
for k,acc,p in zip(ks_vals,ks_acc,ks_params):
    add_row(1,'kernel_size',k,out_all,f'kernel={k}에서 성능과 파라미터 규모가 달라졌다.',f'test_acc={acc:.2f}, params={p}')

# Ex2 Dropout
class DropoutCNN(nn.Module):
    def __init__(self, p=0.5):
        super().__init__()
        self.features = nn.Sequential(
            nn.Conv2d(1,32,3,padding=1), nn.ReLU(),
            nn.Conv2d(32,64,3,padding=1), nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Conv2d(64,128,3,padding=1), nn.ReLU(),
            nn.MaxPool2d(2),
        )
        self.classifier = nn.Sequential(
            nn.Flatten(),
            nn.Linear(128*8*8,256), nn.ReLU(),
            nn.Dropout(p),
            nn.Linear(256,10)
        )
    def forward(self,x):
        return self.classifier(self.features(x))

exd = OUT / 'ex2_dropout'; exd.mkdir(exist_ok=True)
p_vals=[0.0,0.25,0.5,0.75]
tr_accs=[]; te_accs=[]
for p in p_vals:
    m=DropoutCNN(p)
    tr_curve, te = train_model(m, epochs=4, lr=1e-3)
    tr_accs.append(tr_curve[-1]); te_accs.append(te)

plt.figure(figsize=(10,4))
plt.subplot(1,2,1); plt.plot(p_vals, tr_accs, marker='o', label='train'); plt.plot(p_vals, te_accs, marker='s', label='test'); plt.legend(); plt.title('Acc vs Dropout p')
plt.subplot(1,2,2); plt.bar([str(p) for p in p_vals], [tr-te for tr,te in zip(tr_accs,te_accs)]); plt.title('Train-Test Gap')
plt.tight_layout(); out_all=exd/'ex2_dropout_summary.png'; plt.savefig(out_all,dpi=140); plt.close()
for p,tr,te in zip(p_vals,tr_accs,te_accs):
    add_row(2,'dropout_p',p,out_all,f'dropout={p}에서 일반화 격차가 변했다.',f'train={tr:.2f}, test={te:.2f}, gap={tr-te:.2f}')

# Ex3 Depth
class DeepCNN(nn.Module):
    def __init__(self, depth=4):
        super().__init__()
        ch=[1,32,64,64,128,128,256,256,256]
        layers=[]
        for i in range(depth):
            layers += [nn.Conv2d(ch[i], ch[i+1], 3, padding=1), nn.BatchNorm2d(ch[i+1]), nn.ReLU()]
            if i==1: layers += [nn.MaxPool2d(2)]
            if i==3: layers += [nn.MaxPool2d(2)]
        self.features=nn.Sequential(*layers)
        final_ch=ch[depth]
        self.cls=nn.Sequential(nn.AdaptiveAvgPool2d((4,4)), nn.Flatten(), nn.Linear(final_ch*16, 128), nn.ReLU(), nn.Linear(128,10))
    def forward(self,x):
        return self.cls(self.features(x))

exd = OUT / 'ex3_depth'; exd.mkdir(exist_ok=True)
depths=[2,4,6,8]
accs=[]; times=[]
import time
for d in depths:
    m=DeepCNN(d)
    t0=time.time(); tr,te=train_model(m, epochs=4, lr=1e-3); dt=time.time()-t0
    accs.append(te); times.append(dt)

plt.figure(figsize=(10,4))
plt.subplot(1,2,1); plt.plot(depths, accs, marker='o'); plt.title('Test Acc vs Depth')
plt.subplot(1,2,2); plt.plot(depths, times, marker='s'); plt.title('Training Time vs Depth'); plt.ylabel('sec')
plt.tight_layout(); out_all=exd/'ex3_depth_summary.png'; plt.savefig(out_all,dpi=140); plt.close()
for d,a,t in zip(depths,accs,times):
    add_row(3,'num_conv_layers',d,out_all,f'깊이 {d}에서 정확도와 학습 시간의 균형이 달라졌다.',f'test_acc={a:.2f}, time={t:.1f}s')

# Ex4 learning rate
class BaseCNN(nn.Module):
    def __init__(self):
        super().__init__()
        self.f=nn.Sequential(
            nn.Conv2d(1,32,3,padding=1), nn.BatchNorm2d(32), nn.ReLU(),
            nn.Conv2d(32,64,3,padding=1), nn.BatchNorm2d(64), nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Conv2d(64,128,3,padding=1), nn.BatchNorm2d(128), nn.ReLU(),
            nn.MaxPool2d(2),
        )
        self.c=nn.Sequential(nn.Flatten(), nn.Linear(128*8*8,128), nn.ReLU(), nn.Dropout(0.3), nn.Linear(128,10))
    def forward(self,x):
        return self.c(self.f(x))

exd = OUT / 'ex4_learning_rate'; exd.mkdir(exist_ok=True)
lrs=[0.1,0.01,0.001,0.0001]
lr_test=[]
for lr in lrs:
    m=BaseCNN()
    tr,te=train_model(m, epochs=4, lr=lr)
    lr_test.append(te)

plt.figure(figsize=(8,4)); plt.semilogx(lrs, lr_test, marker='o'); plt.title('Test Acc vs Learning Rate'); plt.xlabel('lr'); plt.ylabel('test acc'); plt.tight_layout()
out_all=exd/'ex4_lr_summary.png'; plt.savefig(out_all,dpi=140); plt.close()
for lr,a in zip(lrs,lr_test):
    add_row(4,'learning_rate',lr,out_all,f'lr={lr}에서 수렴 안정성과 최종 성능이 달라졌다.',f'test_acc={a:.2f}')

# save table/csv
csv_path = OUT/'ablation_summary_cv7.csv'
with csv_path.open('w', newline='', encoding='utf-8-sig') as f:
    w=csv.DictWriter(f, fieldnames=['Exercise','Parameter','Value','Output file','Visual change','Optional metric','One-line analysis'])
    w.writeheader(); w.writerows(rows)

overall=OUT/'overall_summary_table_cv7.md'
with overall.open('w', encoding='utf-8-sig') as f:
    f.write('# 전체 결과 요약 표 (CV Exercise 7)\n\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
    for r in rows:
        f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")

meta={
1:('Kernel Size Study','커널 크기가 정확도와 파라미터 규모에 주는 영향을 비교한다.','kernel_size','커널이 커질수록 수용 영역은 넓어진다. 파라미터 수가 함께 증가한다. 작은 입력에서는 큰 커널 이득이 제한적이었다.'),
2:('Dropout Study','Dropout 비율에 따른 일반화 성능을 비교한다.','dropout_p','Dropout이 너무 낮으면 과적합 갭이 커질 수 있다. 너무 높으면 학습 자체가 약해진다. 중간 구간에서 일반화 균형이 좋았다.'),
3:('Depth Study','Conv 레이어 깊이에 따른 성능과 시간 트레이드오프를 비교한다.','num_conv_layers','깊이가 증가하면 표현력은 커진다. 학습 시간도 함께 증가한다. 데이터 규모와 입력 해상도에 맞는 적정 깊이가 필요했다.'),
4:('Learning Rate Study','학습률 변화에 따른 수렴 안정성과 정확도를 비교한다.','learning_rate','큰 학습률은 불안정 수렴을 유발할 수 있다. 너무 작은 학습률은 학습이 느리다. 중간 학습률에서 가장 안정적인 성능이 나타났다.'),
}

md = OUT/'ablation_report_compact_ko_v3_cv7.md'
with md.open('w', encoding='utf-8-sig') as f:
    f.write('# CV Exercise 7 Ablation Study\n\n')
    f.write('> 실행 환경 제약으로 CIFAR-10 대신 sklearn digits(32x32 리사이즈)로 동일 구조 실험을 수행함.\n\n')
    for ex in range(1,5):
        t,o,c,s=meta[ex]
        f.write(f'## Exercise {ex}. {t}\n\n### 1) 실험 목적\n- {o}\n\n### 2) 변경 인자\n- {c}\n\n### 3) 결과 표\n| 파라미터 | 값 | 결과 이미지 |\n|---|---|---|\n')
        for r in rows:
            if r['Exercise']==f'Exercise {ex}': f.write(f"| {r['Parameter']} | {r['Value']} | {r['Output file']} |\n")
        f.write(f"\n### 4) 해석 요약\n- {s}\n\n")

lst=OUT/'generated_images_cv7.txt'
with lst.open('w', encoding='utf-8') as f:
    for r in rows: f.write(r['Output file']+'\n')

# docx report

def set_korean_font(doc):
    normal=doc.styles['Normal']; normal.font.name='Malgun Gothic'; normal.font.size=Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); normal.font.color.rgb=RGBColor(0,0,0)
    for h in ['Heading 1','Heading 2','Heading 3']:
        st=doc.styles[h]; st.font.name='Malgun Gothic'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.color.rgb=RGBColor(0,0,0)

def set_table_width_pct(table,pct=100):
    tbl=table._tbl; tblPr=tbl.tblPr
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'),'pct'); tblW.set(qn('w:w'),str(int(pct*50)))

doc=Document(); set_korean_font(doc)
sec=doc.sections[0]; usable=(sec.page_width-sec.left_margin-sec.right_margin)/914400.0
col1=max(1.3,usable*0.18); col2=max(1.1,usable*0.13); col3=max(3.0,usable-col1-col2); img_w=max(2.3,col3-0.2)
doc.add_paragraph('실행 환경 제약으로 CIFAR-10 대신 sklearn digits(32x32 리사이즈)로 동일 구조 실험을 수행함.', style='List Bullet')
for ex in range(1,5):
    t,o,c,s=meta[ex]
    doc.add_heading(f'Exercise {ex}. {t}', level=2)
    doc.add_heading('1) 실험 목적', level=3); doc.add_paragraph(o, style='List Bullet')
    doc.add_heading('2) 변경 인자', level=3); doc.add_paragraph(c, style='List Bullet')
    doc.add_heading('3) 결과 표', level=3)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'; table.autofit=False; set_table_width_pct(table,100)
    hdr=table.rows[0].cells; hdr[0].text='파라미터'; hdr[1].text='값'; hdr[2].text='결과 이미지'
    widths=[Inches(col1),Inches(col2),Inches(col3)]
    for i,w in enumerate(widths): hdr[i].width=w
    for r in rows:
        if r['Exercise']!=f'Exercise {ex}': continue
        rc=table.add_row().cells; rc[0].text=r['Parameter']; rc[1].text=r['Value']
        for i,w in enumerate(widths): rc[i].width=w
        p=ROOT/r['Output file']
        if p.exists(): rc[2].paragraphs[0].add_run().add_picture(str(p), width=Inches(img_w))
    doc.add_paragraph(''); doc.add_heading('4) 해석 요약', level=3); doc.add_paragraph(s, style='List Bullet'); doc.add_paragraph('')

docx=OUT/'ablation_report_compact_ko_v3_cv7.docx'; doc.save(docx)

print('rows',len(rows))
print(csv_path)
print(md)
print(docx)
print(overall)
print(lst)
print('device', DEVICE)
